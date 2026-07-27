import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_strategy_md_and_docx():
    md_path = r"C:\Users\User\Desktop\TP\outputs\ironman_sub11_and_langkawi_strategy.md"
    docx_path = r"C:\Users\User\Desktop\TP\outputs\ironman_sub11_and_langkawi_strategy.docx"

    lines = []
    lines.append("# Ironman 226km 突破 11 小時 (Sub-11) 與馬來西亞蘭卡威 (Langkawi) 備賽策略指南")
    lines.append("")
    lines.append("本指南針對選手挑戰 **Full Distance Ironman 226km 突破 11 小時 (Sub-11)** 進行時間帳本拆解與訓練缺口分析，並針對 **11 月馬來西亞蘭卡威 (Ironman Malaysia Langkawi)** 的高溫高濕與陡坡地形，提供專屬特化戰術與補給策略。")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## 一、 突破 11 小時 (Sub-11) 的時間帳本與課表差距分析")
    lines.append("")
    lines.append("### 1. Sub-11 小時時間分配帳本 (標準平路/微丘賽道)")
    lines.append("- **🏊 游泳 (3.8km)**：1 小時 12 分 (配速約 1:53/100m)")
    lines.append("- **🚴 T1 轉換**：6 分鐘")
    lines.append("- **🚴 單車 (180km)**：5 小時 35 分 (均速約 32.2 km/h)")
    lines.append("- **🏃 T2 轉換**：5 分鐘")
    lines.append("- **🏃 馬拉松 (42.2km)**：3 小時 55 分 (配速約 5:34 /km)")
    lines.append("- **⏱️ 總完賽時間**：**10 小時 53 分 (突破 11 小時成功！)**")
    lines.append("")
    lines.append("### 2. 突破 11 小時的四大訓練缺口與補強方向")
    lines.append("1. **FTP / 功率體重比儲備不足 (臨門一腳)**：")
    lines.append("   - 現行 FTP = 205W，Zone 2 比賽瓦數為 140W-150W。若體重在 70kg 以上，145W 在一般賽道單車約需 5h 50m ~ 6h 10m，會壓縮馬拉松必須跑出 3h 40m 以內。")
    lines.append("   - **補強建議**：利用備賽期將 **FTP 提升至 220W ~ 235W**（使 Zone 2 賽事瓦數提升至 155W-165W），確保單車穩進 5h 35m 內下車。")
    lines.append("2. **高碳水消化吸收訓練 (High-Carb Fueling Protocol)**：")
    lines.append("   - 必須於週末長騎跑實測推升至 **80g - 100g 碳水化合物/小時**，防止馬拉松 30km 後發生「撞牆」。")
    lines.append("3. **無防寒衣 (Non-Wetsuit) 水感耐力**：")
    lines.append("   - 確保在無防寒衣額外浮力加成下，仍能輕鬆以 1:50-1:55/100m 配速完游 3.8km 且不耗損體力。")
    lines.append("4. **下肢抗疲勞肌力 (Strength & Conditioning)**：")
    lines.append("   - 每週安排 20 分鐘針對臀大肌、臀中肌與核心抗旋轉的單腳肌力訓練。")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## 二、 11 月馬來西亞蘭卡威 (Langkawi) 賽事酷刑與特化策略")
    lines.append("")
    lines.append("### 1. 蘭卡威賽道四大酷刑分析")
    lines.append("| 賽事項目 | 蘭卡威賽道特性 | 對完賽與時間的實際影響 |")
    lines.append("| :--- | :--- | :--- |")
    lines.append("| **氣溫與濕度** | 氣溫 **32°C - 34°C** / 濕度 **85% - 95%** | 高溫高濕酷刑。心率漂移劇烈，馬拉松普遍慢 30-50 分鐘。 |")
    lines.append("| **水溫與游泳** | 水溫 **28°C - 29°C (禁止穿防寒衣)** | 無浮力加成，游泳完賽時間普遍增加 **5-8 分鐘**。 |")
    lines.append("| **單車地形** | **總爬升 1,500m+** (含 15-20% 陡坡) | 粗糙路面與陡坡。FTP 205W 預估單車時間需要 **6h 15m - 6h 40m**。 |")
    lines.append("| **馬拉松賽道** | 機場海岸跑道，**幾乎零遮蔭** | 高 UV 曝曬，耐熱能力決定是跑完還是走完。 |")
    lines.append("")
    lines.append("### 2. 蘭卡威賽事應對總結")
    lines.append("- **備賽時間 (16 週)**：從現在到 11 月中時間完全足夠完成完整 12 週課表！")
    lines.append("- **賽事難度修正**：蘭卡威難度會讓選手比平路賽道慢 **40 - 70 分鐘**。在蘭卡威跑出 11h 30m 的選手，在平路涼爽賽道通常已有 Sub-10:45 實力。")
    lines.append("")
    lines.append("### 3. 蘭卡威三大「專屬特化防禦戰術」")
    lines.append("1. **單車傳動齒比改裝 (Bike Gearing)**：")
    lines.append("   - 面對 15-20% 陡坡，FTP 205W 必須配備 **前 50/34T 壓縮盤 + 後 11-34T (或 36T) 飛輪**，確保上坡踩踏踏頻保持 70+ rpm，避免抽車瓦數飆破 240W 引發抽筋。")
    lines.append("2. **賽前 4-6 週耐熱訓練 (Heat Acclimatization)**：")
    lines.append("   - 週四室內單車關閉風扇/穿長袖騎乘 45 分鐘，或長訓練後進行 **40°C 熱水泡澡 20-30 分鐘**，刺激體內血漿容量擴增。")
    lines.append("3. **冰塊物理降溫與高鈉電解質策略**：")
    lines.append("   - 馬拉松每 2.5km 補給站強制拿冰塊塞帽子與胸口，鹽錠補充提高至 **800mg - 1,200mg 鈉/小時**。")
    lines.append("")

    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"Generated {md_path}")

    # Build docx
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    in_table = False
    table_lines = []

    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        rows = [t.strip("|").split("|") for t in table_lines if t.strip()]
        if len(rows) >= 2:
            if "---" in rows[1][0] or ":---" in rows[1][0]:
                headers = [c.strip() for c in rows[0]]
                data_rows = [[c.strip() for c in r] for r in rows[2:]]
            else:
                headers = [c.strip() for c in rows[0]]
                data_rows = [[c.strip() for c in r] for r in rows[1:]]

            tbl = doc.add_table(rows=len(data_rows)+1, cols=len(headers))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            hdr_cells = tbl.rows[0].cells
            for idx, text in enumerate(headers):
                hdr_cells[idx].text = text
                set_cell_background(hdr_cells[idx], "1A365D")
                p = hdr_cells[idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.name = "微軟正黑體"
                    run.font.size = Pt(10)

            for r_idx, r_data in enumerate(data_rows):
                row_cells = tbl.rows[r_idx+1].cells
                bg_color = "F7FAFC" if r_idx % 2 == 1 else "FFFFFF"
                for c_idx, val in enumerate(r_data):
                    if c_idx < len(row_cells):
                        row_cells[c_idx].text = val
                        set_cell_background(row_cells[c_idx], bg_color)
                        p = row_cells[c_idx].paragraphs[0]
                        for run in p.runs:
                            run.font.name = "微軟正黑體"
                            run.font.size = Pt(9.5)
                        if c_idx == 0:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        table_lines = []

    for line in lines:
        sline = line.strip()
        if sline.startswith("|"):
            in_table = True
            table_lines.append(sline)
            continue
        elif in_table:
            in_table = False
            flush_table()

        if sline.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(sline[2:].strip())
            run.font.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(26, 54, 93)
            run.font.name = "微軟正黑體"
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
        elif sline.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(sline[3:].strip())
            run.font.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(43, 108, 176)
            run.font.name = "微軟正黑體"
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
        elif sline.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(sline[4:].strip())
            run.font.bold = True
            run.font.size = Pt(12.5)
            run.font.color.rgb = RGBColor(45, 55, 72)
            run.font.name = "微軟正黑體"
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif sline.startswith("- ") or sline.startswith("1. ") or sline.startswith("2. ") or sline.startswith("3. ") or sline.startswith("4. "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            parts = sline.split("**")
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = "微軟正黑體"
                run.font.size = Pt(10.5)
                if idx % 2 == 1:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(26, 54, 93)
            p.paragraph_format.space_after = Pt(3)
        elif sline.startswith("   - "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            parts = sline.strip().split("**")
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = "微軟正黑體"
                run.font.size = Pt(10)
                if idx % 2 == 1:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(43, 108, 176)
            p.paragraph_format.space_after = Pt(2)
        elif sline == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        elif sline:
            p = doc.add_paragraph()
            parts = sline.split("**")
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = "微軟正黑體"
                run.font.size = Pt(10.5)
                if idx % 2 == 1:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(26, 54, 93)
            p.paragraph_format.space_after = Pt(4)

    if in_table:
        flush_table()

    doc.save(docx_path)
    print(f"Generated {docx_path}")

def create_ftp_md_and_docx():
    md_path = r"C:\Users\User\Desktop\TP\outputs\ftp_elevation_guide_and_workouts.md"
    docx_path = r"C:\Users\User\Desktop\TP\outputs\ftp_elevation_guide_and_workouts.docx"

    lines = []
    lines.append("# FTP 高效提升專項指南與課表 (Based on FTP 205W)")
    lines.append("")
    lines.append("本指南專為 **現行 FTP = 205 W** 的鐵人三項選手設計。深入剖析提升 FTP 的生理機制（VO2max 天花板與乳酸清除能力），並提供量化功率課表、4 週輪替循環與突破關鍵輔助策略。")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## 一、 FTP 提升的核心生理機制")
    lines.append("FTP (Functional Threshold Power) 代表在 1 小時內能維持的最大平均功率，由兩個核心要素決定：")
    lines.append("1. **最大攝氧量 (VO2max — 引擎天花板)**：代表最大吸氧能力。天花板越高，潛力越大。")
    lines.append("2. **乳酸閾值百分比 (% of VO2max — 天花板利用率)**：代表肌肉能在多高強度下「邊高速踩踏、邊清除乳酸」。")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## 二、 四大高效率 FTP 提升專項課表 (FTP 205W 量化)")
    lines.append("")
    lines.append("### 1. Sweet Spot (甜心區間：88% - 94% FTP) — 最經濟高效瓦數累積")
    lines.append("- **目標瓦數 (FTP 205W)**：**180 W - 193 W** (目標定在 **185 W**)")
    lines.append("- **特點**：強度剛好低於乳酸閾值，疲勞累積慢、恢復快，能最大化刺激線粒體增生與毛細血管密化。")
    lines.append("- **菜單**：15m 熱身 ➔ **3 x 15 分鐘 @ 185W** (組間休息 4 分鐘 110W 輕鬆踩) ➔ 15m 緩冷。")
    lines.append("")
    lines.append("### 2. Threshold (乳酸閾值間歇：95% - 105% FTP) — 直接攻克 FTP 本體")
    lines.append("- **目標瓦數 (FTP 205W)**：**195 W - 215 W** (目標定在 **205 W**)")
    lines.append("- **特點**：直接在 FTP 本體強度踩踏，強化心理耐受力與乳酸清除效率。")
    lines.append("- **菜單**：15m 熱身 ➔ **4 x 8 分鐘 @ 205W** (組間休息 4 分鐘 110W 輕鬆踩) ➔ 15m 緩冷。")
    lines.append("")
    lines.append("### 3. Over-Under (過衝/拉回間歇) — 訓練高速清除乳酸能力")
    lines.append("- **目標瓦數 (FTP 205W)**：Over **215 W (105% FTP)** / Under **185 W (90% FTP)**")
    lines.append("- **特點**：在 215W 產生乳酸，接著降至 185W 強迫肌肉在高速運轉中將乳酸轉化為能量。")
    lines.append("- **菜單**：15m 熱身 ➔ **3 組 x 12 分鐘 Over-Under** (每組包含 4 次：[2m @ 215W + 1m @ 185W]，組間休息 5m 110W) ➔ 15m 緩冷。")
    lines.append("")
    lines.append("### 4. VO2max (最大攝氧量間歇：108% - 120% FTP) — 打開瓦數天花板")
    lines.append("- **目標瓦數 (FTP 205W)**：**220 W - 245 W** (目標定在 **230 W**)")
    lines.append("- **特點**：極短時間高瓦數刺激，把最大吸氧能力衝高，打破 FTP 停滯瓶頸。")
    lines.append("- **菜單**：15m 熱身 ➔ **5 x 3 分鐘 @ 230W** (迴轉數 95-100 rpm，組間休息 3m 110W) ➔ 20m 緩冷。")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## 三、 融入每週單車課表的 4 週進階輪替計畫")
    lines.append("")
    lines.append("| 週次 | 週四單車主課表內容 (FTP 205W 專屬) | 騎後跑 (Transition Run) |")
    lines.append("| :---: | :--- | :--- |")
    lines.append("| **第 1 週** | **Sweet Spot 基礎組**：15m 熱身 ➔ **3 x 15m @ 185W** (休息 4m) ➔ 15m 緩冷 | **10m 轉接跑** (Zone 2, 90rpm) |")
    lines.append("| **第 2 週** | **Threshold 閾值攻堅**：15m 熱身 ➔ **4 x 8m @ 205W** (休息 4m) ➔ 15m 緩冷 | **15m 轉接跑** (Zone 2, 90rpm) |")
    lines.append("| **第 3 週** | **Over-Under 清除組**：15m 熱身 ➔ **3 x 12m (215W/185W 輪替)** (休息 5m) ➔ 15m 緩冷 | **15m 轉接跑** (Zone 2, 90rpm) |")
    lines.append("| **第 4 週** | **VO2max 天花板衝刺**：15m 熱身 ➔ **5 x 3m @ 230W** (休息 3m) ➔ 20m 輕鬆 | **10m 轉接跑** (輕鬆慢跑) |")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## 四、 FTP 突破的四大關鍵輔助策略")
    lines.append("1. **高踏頻踩踏 (High Cadence Strategy)**：間歇維持 **90 - 98 rpm**，減輕肌肉疲勞並將負荷轉移給心血管。")
    lines.append("2. **訓練後 30 分鐘黃金補給**：補充 3:1 或 4:1 的碳水與蛋白質（如 60g 碳水 + 20g 蛋白質），加速糖原合成與肌肉修復。")
    lines.append("3. **騎行台 ERG 模式訓練**：強迫雙腿輸出精確瓦數，不給降瓦偷懶空間。")
    lines.append("4. **4-6 週重新測試 FTP (Ramp Test)**：定期重新測驗並更新功率區間，持續推升訓練刺激。")
    lines.append("")

    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"Generated {md_path}")

    # Build docx
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    in_table = False
    table_lines = []

    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        rows = [t.strip("|").split("|") for t in table_lines if t.strip()]
        if len(rows) >= 2:
            if "---" in rows[1][0] or ":---" in rows[1][0]:
                headers = [c.strip() for c in rows[0]]
                data_rows = [[c.strip() for c in r] for r in rows[2:]]
            else:
                headers = [c.strip() for c in rows[0]]
                data_rows = [[c.strip() for c in r] for r in rows[1:]]

            tbl = doc.add_table(rows=len(data_rows)+1, cols=len(headers))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            hdr_cells = tbl.rows[0].cells
            for idx, text in enumerate(headers):
                hdr_cells[idx].text = text
                set_cell_background(hdr_cells[idx], "1A365D")
                p = hdr_cells[idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.name = "微軟正黑體"
                    run.font.size = Pt(10)

            for r_idx, r_data in enumerate(data_rows):
                row_cells = tbl.rows[r_idx+1].cells
                bg_color = "F7FAFC" if r_idx % 2 == 1 else "FFFFFF"
                for c_idx, val in enumerate(r_data):
                    if c_idx < len(row_cells):
                        row_cells[c_idx].text = val
                        set_cell_background(row_cells[c_idx], bg_color)
                        p = row_cells[c_idx].paragraphs[0]
                        for run in p.runs:
                            run.font.name = "微軟正黑體"
                            run.font.size = Pt(9.5)
                        if c_idx == 0:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        table_lines = []

    for line in lines:
        sline = line.strip()
        if sline.startswith("|"):
            in_table = True
            table_lines.append(sline)
            continue
        elif in_table:
            in_table = False
            flush_table()

        if sline.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(sline[2:].strip())
            run.font.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(26, 54, 93)
            run.font.name = "微軟正黑體"
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
        elif sline.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(sline[3:].strip())
            run.font.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(43, 108, 176)
            run.font.name = "微軟正黑體"
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
        elif sline.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(sline[4:].strip())
            run.font.bold = True
            run.font.size = Pt(12.5)
            run.font.color.rgb = RGBColor(45, 55, 72)
            run.font.name = "微軟正黑體"
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif sline.startswith("- ") or sline.startswith("1. ") or sline.startswith("2. ") or sline.startswith("3. ") or sline.startswith("4. "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            parts = sline.split("**")
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = "微軟正黑體"
                run.font.size = Pt(10.5)
                if idx % 2 == 1:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(26, 54, 93)
            p.paragraph_format.space_after = Pt(3)
        elif sline == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        elif sline:
            p = doc.add_paragraph()
            parts = sline.split("**")
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = "微軟正黑體"
                run.font.size = Pt(10.5)
                if idx % 2 == 1:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(26, 54, 93)
            p.paragraph_format.space_after = Pt(4)

    if in_table:
        flush_table()

    doc.save(docx_path)
    print(f"Generated {docx_path}")

if __name__ == "__main__":
    create_strategy_md_and_docx()
    create_ftp_md_and_docx()
