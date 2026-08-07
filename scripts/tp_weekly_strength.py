from __future__ import annotations

import argparse
import re
import json
import urllib.request
from datetime import date, datetime, timedelta
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import sys
ROOT = Path(r"C:\Users\User\Desktop\TP")
if str(ROOT / "scripts") not in sys.path:
    sys.path.append(str(ROOT / "scripts"))
from estimator_226 import calculate_dynamic_226_estimate

CONFIG = ROOT / "config" / "trainingpeaks_calendar_url.txt"
OUT_DIR = ROOT / "outputs" / "weekly"
EXCEL_PATH = ROOT / "運動.xlsx"
CACHE_FILE = ROOT / "data" / "raw" / "calendar_cache.json"
WEEKDAYS = "一二三四五六日"


def font(run, size=11, bold=False, color="000000"):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def shade(cell, fill):
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(node)


def table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr
    tblw = tblpr.first_child_found_in("w:tblW")
    tblw.set(qn("w:w"), str(sum(widths)))
    tblw.set(qn("w:type"), "dxa")
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    tblpr.append(ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(width))
        grid.append(node)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.first_child_found_in("w:tcW")
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            margins = OxmlElement("w:tcMar")
            for name, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
                node = OxmlElement(f"w:{name}")
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")
                margins.append(node)
            tcpr.append(margins)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def exercise_table(doc, title, rows):
    doc.add_paragraph(title, style="Heading 2")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "動作"
    table.rows[0].cells[1].text = "建議量"
    for cell in table.rows[0].cells:
        shade(cell, "E8EEF5")
    for exercise, dose in rows:
        cells = table.add_row().cells
        cells[0].text, cells[1].text = exercise, dose
    table_geometry(table, [3600, 5760])
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    font(run, 10.5, row_index == 0)


def label_from_cached_events(events):
    if not events:
        return "未排定"
    labels = []
    for ev in events:
        text = ev.get("summary", "未命名訓練")
        text = text.replace("Day Off:", "休息：")
        labels.append(text)
    return "＋".join(labels)


def allocate_from_cache(this_week_dict, monday):
    labels = {day: label_from_cached_events(events) for day, events in this_week_dict.items()}
    plan = {day: "不排肌力；專注完成 TP 主課。" for day in this_week_dict}
    main = monday
    if not re.search(r"休息|Day Off|未排定", labels[main], re.I):
        candidates = [
            day for day in sorted(labels.keys())
            if not re.search(r"長|4:30|90 k|TEMPO|間歇|Tempo", labels[day], re.I)
        ]
        if candidates:
            main = min(candidates, key=lambda day: len(labels[day]))
    variant = "A" if monday.isocalendar().week % 2 else "B"
    plan[main] = f"主課 {variant}，30–40 分鐘；保留 2–3 下餘裕，不做到力竭。"
    short_count = 0
    for day in sorted(labels.keys()):
        if day == main:
            continue
        text = labels[day]
        if re.search(r"4:30|長|90 k|大安", text, re.I):
            plan[day] = "不排肌力；長距離日僅做 8–10 分鐘活動度。"
        elif "Swim" in text and short_count < 2:
            plan[day] = "短課 15–20 分鐘；游泳後 30–60 分鐘或晚間完成。"
            short_count += 1
    return plan, variant


def build(monday):
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    sunday = monday + timedelta(days=6)
    last_monday = monday - timedelta(days=7)
    last_sunday = last_monday + timedelta(days=6)
    
    # Load cache
    cache = {}
    if CACHE_FILE.exists():
        try:
            with open(CACHE_FILE, "r", encoding="utf-8") as f:
                cache = json.load(f)
        except Exception as e:
            print("Error loading cache:", e)
            
    # Group cached events by date
    cached_by_date = {}
    curr_date = last_monday
    while curr_date <= sunday:
        cached_by_date[curr_date] = []
        curr_date += timedelta(days=1)
        
    for uid, ev in cache.items():
        ev_date = date.fromisoformat(ev["date"])
        if ev_date in cached_by_date:
            cached_by_date[ev_date].append(ev)
            
    # 1. Analyze last week data (planned vs actual)
    last_week_events = []
    for d in sorted(cached_by_date.keys()):
        if last_monday <= d <= last_sunday:
            for ev in cached_by_date[d]:
                orig = ev.get("original_plan", {})
                planned_time = orig.get("planned_time", ev.get("planned_time", 0))
                planned_dist = orig.get("planned_dist", ev.get("planned_dist", 0))
                
                last_week_events.append({
                    "type": ev.get("type", "Other"),
                    "planned_time": planned_time,
                    "actual_time": ev.get("actual_time", 0),
                    "planned_dist": planned_dist,
                    "actual_dist": ev.get("actual_dist", 0)
                })
                
    df_last = pd.DataFrame(last_week_events) if last_week_events else pd.DataFrame(columns=["type", "planned_time", "actual_time", "planned_dist", "actual_dist"])
    
    # Calculate completion rates
    last_week_stats = {}
    for t in ["Swim", "Bike", "Run", "Strength"]:
        sub = df_last[df_last["type"] == t] if not df_last.empty else pd.DataFrame()
        pt = sub["planned_time"].sum() if not sub.empty else 0
        at = sub["actual_time"].sum() if not sub.empty else 0
        pd_sum = sub["planned_dist"].sum() if not sub.empty else 0
        ad_sum = sub["actual_dist"].sum() if not sub.empty else 0
        last_week_stats[t] = {
            "planned_time": pt / 60.0,
            "actual_time": at / 60.0,
            "pct_time": (at / pt * 100) if pt > 0 else 0,
            "planned_dist": pd_sum,
            "actual_dist": ad_sum,
            "pct_dist": (ad_sum / pd_sum * 100) if pd_sum > 0 else 0,
            "sessions": len(sub) if not sub.empty else 0
        }
        
    total_pt = df_last["planned_time"].sum() if not df_last.empty else 0
    total_at = df_last["actual_time"].sum() if not df_last.empty else 0
    overall_completion = (total_at / total_pt * 100) if total_pt > 0 else 0
    
    # 2. Load 226 race history
    races_226 = []
    if EXCEL_PATH.exists():
        try:
            df_tri = pd.read_excel(EXCEL_PATH, sheet_name="鐵人賽")
            df_226 = df_tri[df_tri["距離"] == 226].dropna(subset=["總時間"])
            for _, r in df_226.iterrows():
                races_226.append({
                    "event": str(r.get("賽事", "未知賽事")),
                    "swim": str(r.get("游泳", "-")),
                    "t1": str(r.get("T1", "-")),
                    "bike": str(r.get("騎車", "-")),
                    "t2": str(r.get("T2", "-")),
                    "run": str(r.get("跑步", "-")),
                    "total": str(r.get("總時間", "-"))
                })
        except Exception as e:
            print("Error loading Excel:", e)

    # 3. Strength allocation for this week
    this_week_dict = {d: cached_by_date[d] for d in sorted(cached_by_date.keys()) if monday <= d <= sunday}
    plans, variant = allocate_from_cache(this_week_dict, monday)
    
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = "Calibri", Pt(11)
    normal.paragraph_format.space_after, normal.paragraph_format.line_spacing = Pt(6), 1.25
    
    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, "1F4D78"),
        ("Heading 2", 13, 14, 7, "2E74B5"),
        ("Heading 3", 11.5, 10, 5, "1F4D78"),
    ):
        style = doc.styles[name]
        style.font.name, style.font.size = "Calibri", Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before, style.paragraph_format.space_after = Pt(before), Pt(after)
        
    header = section.header.paragraphs[0]
    header.text = "TrainingPeaks 週訓練整合與完賽成效分析"
    font(header.runs[0], 9, False, "666666")
    
    footer = section.footer.paragraphs[0]
    footer.text = f"{monday:%Y-%m-%d} 至 {sunday:%Y-%m-%d} | 本週主課 {variant}"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(footer.runs[0], 9, False, "777777")
    
    # Document Title
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    font(title.add_run(f"第 {monday.isocalendar().week:02d} 週訓練整合、完賽分析與肌力課表"), 22, True, "1F4D78")
    
    meta = doc.add_paragraph()
    font(meta.add_run(f"週期：{monday:%Y/%m/%d}－{sunday:%Y/%m/%d}　｜　來源：TrainingPeaks & 運動.xlsx"), 10, False, "555555")
    
    # ------------------ SECTION 1 ------------------
    doc.add_paragraph("一、 上週訓練執行率回顧", style="Heading 1")
    lead1 = doc.add_paragraph()
    font(lead1.add_run(f"上週（{last_monday:%m/%d}–{last_sunday:%m/%d}）整體課表時間完成度達 "), 11)
    font(lead1.add_run(f"{overall_completion:.1f}%"), 11, True, "C00000")
    font(lead1.add_run("，有氧基礎扎實。詳細完成率數據如下表："), 11)
    
    table1 = doc.add_table(rows=1, cols=7)
    table1.style = "Table Grid"
    headers1 = ["運動項目", "計劃時間", "實際時間", "時間執行率", "計劃距離", "實際距離", "距離執行率"]
    for idx, text in enumerate(headers1):
        table1.rows[0].cells[idx].text = text
        shade(table1.rows[0].cells[idx], "1F4D78")
        for p in table1.rows[0].cells[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                font(run, 10, True, "FFFFFF")
                
    for t in ["Swim", "Bike", "Run", "Strength"]:
        stats = last_week_stats[t]
        cells = table1.add_row().cells
        cells[0].text = "游泳 (Swim)" if t == "Swim" else "單車 (Bike)" if t == "Bike" else "跑步 (Run)" if t == "Run" else "肌力 (Strength)"
        cells[1].text = f"{stats['planned_time']:.2f} 小時"
        cells[2].text = f"{stats['actual_time']:.2f} 小時"
        cells[3].text = f"{stats['pct_time']:.1f}%"
        cells[4].text = f"{stats['planned_dist']:.2f} km"
        cells[5].text = f"{stats['actual_dist']:.2f} km"
        cells[6].text = f"{stats['pct_dist']:.1f}%" if t != "Strength" else "-"
        
    table_geometry(table1, [1500, 1300, 1300, 1300, 1300, 1300, 1360])
    for row_idx, row in enumerate(table1.rows):
        if row_idx == 0: continue
        for col_idx, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    font(run, 9.5, False)
                    
    # Analysis Commentary
    analysis_p = doc.add_paragraph()
    analysis_p.paragraph_format.space_before = Pt(8)
    font(analysis_p.add_run("💡 執行率分析與建議：\n"), 11, True, "1F4D78")
    
    # Check if data was truncated (if Monday/Tuesday of last week is not in cache)
    has_mon_tue = False
    for uid, ev in cache.items():
        if ev["date"] in {str(last_monday), str(last_monday + timedelta(days=1))}:
            has_mon_tue = True
            break
            
    if not has_mon_tue:
        font(analysis_p.add_run("⚠️ 提醒：由於今日剛啟用「每日行事曆快取系統」，上週一部分（週一/週二）的歷史數據在 TrainingPeaks Webcal 串流中已失效，本次執行率可能偏低。從下週一開始，系統將透過每日快取產出完整無缺的計畫執行率對比報告！\n"), 10, True, "9C3D10")
        
    swim_s = last_week_stats['Swim']
    bike_s = last_week_stats['Bike']
    run_s = last_week_stats['Run']
    str_s = last_week_stats['Strength']

    swim_txt = f"1. 游泳狀況：上週游泳共完成 {swim_s['sessions']} 次（累計 {swim_s['actual_dist']:.2f} km / {swim_s['actual_time']:.2f} 小時），達計劃時間之 {swim_s['pct_time']:.1f}%。這有助於奠定強大的水域自信，在 226 比賽中能以低能耗起水。\n" if swim_s['sessions'] > 0 else "1. 游泳狀況：上週未紀錄游泳訓練，建議本週維持基本水感練習。\n"
    bike_txt = f"2. 單車狀況：上週單車共完成 {bike_s['sessions']} 次（累計 {bike_s['actual_dist']:.2f} km / {bike_s['actual_time']:.2f} 小時），達計劃時間之 {bike_s['pct_time']:.1f}%。有氧耐力與瓦數輸出保持穩定。\n" if bike_s['sessions'] > 0 else "2. 單車狀況：上週未紀錄單車騎乘，建議本週恢復長騎課表以維持功率與耐力。\n"
    run_txt = f"3. 跑步狀況：上週跑步共完成 {run_s['sessions']} 次（累計 {run_s['actual_dist']:.2f} km / {run_s['actual_time']:.2f} 小時），達計劃時間之 {run_s['pct_time']:.1f}%。轉換跑與長跑步頻節奏保持良好，請繼續維持肌肉衝擊適應。\n" if run_s['sessions'] > 0 else "3. 跑步狀況：上週未紀錄跑步訓練，本週請注意漸進恢復跑量與步頻感受。\n"
    str_txt = f"4. 肌力狀況：上週肌力共完成 {str_s['sessions']} 次（累計 {str_s['actual_time']:.2f} 小時），達計劃時間之 {str_s['pct_time']:.1f}%。核心與下肢單腿穩定性維持良好，有效支援騎跑表現。" if str_s['sessions'] > 0 else "4. 肌力狀況：上週未紀錄肌力訓練，本週請按課表執行主課與短課，以預防運動傷害。"

    font(analysis_p.add_run(swim_txt))
    font(analysis_p.add_run(bike_txt))
    font(analysis_p.add_run(run_txt))
    font(analysis_p.add_run(str_txt))
    
    # ------------------ SECTION 2 ------------------
    doc.add_paragraph("二、 226 完賽成績分析與預估", style="Heading 1")
    doc.add_paragraph("您的歷史 226 完賽數據總覽：")
    
    if races_226:
        table2 = doc.add_table(rows=1, cols=7)
        table2.style = "Table Grid"
        headers2 = ["賽事名稱", "游泳", "T1", "自行車", "T2", "跑步", "總時間"]
        for idx, text in enumerate(headers2):
            table2.rows[0].cells[idx].text = text
            shade(table2.rows[0].cells[idx], "595959")
            for p in table2.rows[0].cells[idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    font(run, 10, True, "FFFFFF")
                    
        for race in races_226:
            cells = table2.add_row().cells
            cells[0].text = race["event"]
            cells[1].text = race["swim"]
            cells[2].text = race["t1"]
            cells[3].text = race["bike"]
            cells[4].text = race["t2"]
            cells[5].text = race["run"]
            cells[6].text = race["total"]
            
        table_geometry(table2, [2600, 1100, 900, 1200, 900, 1100, 1560])
        for row_idx, row in enumerate(table2.rows):
            if row_idx == 0: continue
            for col_idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        font(run, 9, False)
                        
    # Predict Commentary with Dynamic Estimation Model
    est = calculate_dynamic_226_estimate(monday)
    bench = est["benchmark_comparison"]
    forms = est["formulas"]

    predict_p = doc.add_paragraph()
    predict_p.paragraph_format.space_before = Pt(10)
    font(predict_p.add_run("🔮 IM226 滾動動態完賽成績預估模型 (基於近 4 週訓練與公開標竿數據)：\n"), 11, True, "1F4D78")
    font(predict_p.add_run(f"• 歷史基準：個人 PB 11:38:24 (2024 普悠瑪) | 目前 FTP: {bench['cohort_ftp']}\n"))
    font(predict_p.add_run(f"• 近 4 週動態指標：滾動平均週跑量 {est['rolling_4w_avg_run_km']} km、週騎量 {est['rolling_4w_avg_bike_km']} km、週游量 {est['rolling_4w_avg_swim_km']} km (整體執行率 {est['rolling_4w_exec_rate']}%)\n"))
    font(predict_p.add_run("• 經典公開公式與演算表交叉比對：\n"))
    font(predict_p.add_run(f"  - Riegel 體力疲勞指數公式 (T2=T1*(D2/D1)^1.06): {forms['riegel_power_law']}\n"))
    font(predict_p.add_run(f"  - 70.3 轉 226 經驗倍率法 (2.12x ~ 2.28x): {forms['multiplier_703']}\n"))
    font(predict_p.add_run(f"  - Alan Couzens / Joe Friel 功率能耗模型: {forms['couzens_pacing']}\n"))
    font(predict_p.add_run("• 本週最新動態預估時間（依據訓練完成度滾動計算）：\n"))
    font(predict_p.add_run(f"  - 樂觀目標 (高峰發揮 / 全場未抽筋)：{est['optimistic_range']} (預估中位數 {est['optimistic_mid']})\n"), 11, True, "385623")
    font(predict_p.add_run(f"  - 中性目標 (穩定完賽 / 體力適當分配)：{est['neutral_range']} (預估中位數 {est['neutral_mid']})\n"), 11, True, "C55A11")
    font(predict_p.add_run(f"  - 保守目標 (後半程馬拉松掉速/抽筋)：{est['conservative_range']} (預估中位數 {est['conservative_mid']})\n"), 11, True, "C00000")
    
    # ------------------ SECTION 3 ------------------
    doc.add_paragraph("三、 本週訓練安排與肌力整合", style="Heading 1")
    doc.add_paragraph("本週 TrainingPeaks 行事曆課表與肌力訓練整合表：")
    
    table3 = doc.add_table(rows=1, cols=3)
    table3.style = "Table Grid"
    for cell, text in zip(table3.rows[0].cells, ("日期", "TrainingPeaks 課表", "肌力課表安排")):
        cell.text = text
        shade(cell, "E8EEF5")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                font(run, 10, True, "1F4D78")
                
    for day in sorted(this_week_dict.keys()):
        cells = table3.add_row().cells
        cells[0].text = f"{day:%m/%d}\n週{WEEKDAYS[day.weekday()]}"
        cells[1].text = label_from_cached_events(this_week_dict[day])
        cells[2].text = plans.get(day, "無安排")
        
    table_geometry(table3, [1350, 3600, 4410])
    for row_index, row in enumerate(table3.rows):
        for col_index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_index == 0 or row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    font(run, 9.5, row_index == 0)
                    
    # ------------------ SECTION 4 ------------------
    iso_week = monday.isocalendar().week
    if iso_week <= 30:
        phase_num = 1
        phase_title = "第一階段：解剖適應與防傷基礎期 (Weeks 27–30)"
        main_a = [
            ("分腿蹲 (Split Squat)", "3 組 × 6–8 下／邊 (強化單腿力量與平衡)"),
            ("羅馬尼亞硬舉 (RDL)", "3 組 × 6–8 下 (強化臀大肌與大腿後側膕繩肌)"),
            ("單腳硬舉 (Single Leg RDL)", "2 組 × 6–8 下／邊 (訓練腳踝與髖關節穩定)"),
            ("側棒式 (Side Plank)", "2 組 × 30–45 秒／邊 (強化核心側向抗側彎能力)"),
            ("死蟲式 (Dead bug)", "2 組 × 8–10 下／邊 (核心抗伸展能力)"),
            ("提踵 (Calf Raises)", "3 組 × 12–15 下 (加強小腿與跟腱強度)")
        ]
        main_b = [
            ("臀推或橋式 (Hip Thrust)", "3 組 × 8–10 下 (活化並強化臀大肌)"),
            ("登階 (Step-up)", "3 組 × 6–8 下／邊 (模仿跑步推蹬與單腳穩定)"),
            ("帕羅夫壓 (Pallof press)", "2 組 × 10–12 下／邊 (核心抗旋轉能力)"),
            ("怪獸走 (Monster Walk)", "2 組 × 10–15 步／邊 (激活臀中肌，防止膝外翻)"),
            ("離心提踵 (Eccentric Calf)", "3 組 × 8–12 下 (預防跟腱炎，加強推蹬剛性)"),
            ("髖屈肌伸展 (Hip Flexor Stretch)", "2 組 × 30 秒／邊 (舒緩久坐與騎車的緊繃)")
        ]
        short = [
            ("死蟲式 (Dead bug)", "2 組 × 8–10 下／邊"),
            ("側棒式 (Side Plank)", "2 組 × 30 秒／邊"),
            ("鳥狗式 (Bird dog)", "2 組 × 8 下／邊 (強化對側核心控制)"),
            ("彈力帶側走 (Band Walk)", "2 組 × 10–15 步／邊"),
            ("單腳提踵 (Single Calf Raise)", "2 組 × 12–15 下／邊")
        ]
    elif iso_week <= 34:
        phase_num = 2
        phase_title = "第二階段：力量進階與單腿爆發期 (Weeks 31–34)"
        main_a = [
            ("保加利亞分腿蹲 (Bulgarian Split Squat)", "3 組 × 6–8 下／邊 (大幅提升單腿負重與髖關節伸展)"),
            ("負重羅馬尼亞硬舉 (Barbell/DB RDL)", "3 組 × 6 下 (強化後側鏈拉力與下背保護)"),
            ("哥薩克深蹲 (Cossack Squat)", "2 組 × 6–8 下／邊 (強化內收肌與髖關節多維度活動度)"),
            ("側棒式抬腿 (Side Plank w/ Leg Lift)", "2 組 × 8–10 下／邊 (動態強化臀中肌與抗側彎)"),
            ("抗拉力帶鳥狗式 (Resisted Bird-Dog)", "2組 × 8–10 下／邊 (加強對側核心傳導)"),
            ("單腳離心提踵 (Single Calf Drop)", "3 組 × 10–12 下／邊 (提升跟腱剛性與跑步推蹬耐受)")
        ]
        main_b = [
            ("單腿臀推 (Single-leg Hip Thrust)", "3 組 × 8–10 下／邊 (活化臀大肌單側獨立輸出)"),
            ("負重登階加提膝 (Weighted Step-up w/ Knee Drive)", "3 組 × 6–8 下／邊 (完全模擬跑步推蹬與擺腿機制)"),
            ("壺鈴搖擺 (Kettlebell Swing)", "3 組 × 12–15 下 (訓練髖關節快速伸展爆發力)"),
            ("站姿帕羅夫壓帶轉體 (Pallof Press w/ Rotation)", "2 組 × 10 下／邊 (動態抗旋轉與軀幹穩定)"),
            ("懸吊/滾輪核心抗伸展 (Ab Wheel Rollout)", "2 組 × 8–10 下 (強化前側腹直肌與深層核心)"),
            ("站姿彈力帶提膝 (Band Hip Flexor Drive)", "2 組 × 12 下／邊 (強化長跑擺腿抗疲勞能力)")
        ]
        short = [
            ("滾輪核心抗伸展 (Ab Wheel Rollout)", "2 組 × 8 下"),
            ("側棒式轉體 (Side Plank Threading)", "2 組 × 8 下／邊"),
            ("單腿硬舉靜態支撐 (Single Leg RDL Hold)", "2 組 × 20 秒／邊"),
            ("彈力帶怪獸走+深蹲 (Band Walk to Squat)", "2 組 × 12 步"),
            ("負重單腳提踵 (Weighted Single Calf Raise)", "2 組 × 10–12 下／邊")
        ]
    elif iso_week <= 38:
        phase_num = 3
        phase_title = "第三階段：鐵人專項耐力與神經傳導期 (Weeks 35–38)"
        main_a = [
            ("跳躍登階 / 爆發分腿蹲 (Plyometric Split Jump)", "3 組 × 5 下／邊 (提升神經快縮與跑姿彈性)"),
            ("高次數羅馬尼亞硬舉 (RDL)", "3 組 × 10–12 下 (強化肌耐力與跑後段抗衰退能力)"),
            ("單腳硬舉加提膝 (Single Leg RDL to Knee Drive)", "3 組 × 6–8 下／邊 (整合跑步全週期連結)"),
            ("側棒式轉體抬腿 (Dynamic Side Plank Threading)", "2 組 × 8 下／邊 (極限抗側彎與動態穩定)"),
            ("死蟲式加抗力球 (Deadbug w/ Stability Ball)", "2 組 × 10 下／邊 (強化核心協調性)"),
            ("提踵加跳躍 (Calf Pogo Jumps)", "3 組 × 15 秒 (提升腳踝跟腱彈性反應)")
        ]
        main_b = [
            ("單腿負重臀推 (Weighted Single Hip Thrust)", "3 組 × 8 下／邊 (建立強大爬坡與騎車推蹬功率)"),
            ("高箱登階 (High Step-Up)", "3 組 × 8 下／邊 (加深髖關節屈曲肌群刺激)"),
            ("地雷管旋轉 / 帕羅夫旋轉 (Landmine Rotation)", "2 組 × 10 下／邊 (轉化核心剛性至跑步擺臂)"),
            ("彈力帶橫向快速走 (Speed Band Walk)", "2 組 × 15 步／邊 (維繫髖關節外展肌群高頻率)"),
            ("離心提踵 (Eccentric Calf)", "3 組 × 10 下 (強化跟腱韌性)"),
            ("髖屈肌拉伸與活化 (Hip Flexor Stretch & Drive)", "2 組 × 30 秒／邊")
        ]
        short = [
            ("死蟲式 (Dead bug)", "2 組 × 10 下／邊"),
            ("側棒式 (Side Plank)", "2 組 × 40 秒／邊"),
            ("鳥狗式 (Bird dog)", "2 組 × 10 下／邊"),
            ("彈力帶側走 (Band Walk)", "2 組 × 15 步／邊"),
            ("單腳提踵 (Single Calf Raise)", "2 組 × 15 下／邊")
        ]
    else:
        phase_num = 4
        phase_title = "第四階段：賽前減量與神經活化期 (Weeks 39+)"
        main_a = [
            ("輕量分腿蹲 (Light Split Squat)", "2 組 × 5 下／邊 (神經活化不累積疲勞)"),
            ("輕量單腳硬舉 (Light Single Leg RDL)", "2 組 × 5 下／邊 (本體感覺維繫)"),
            ("側棒式 (Side Plank)", "2 組 × 20 秒／邊 (核心維力)"),
            ("提踵 (Calf Raises)", "2 組 × 10 下 (腳踝喚醒)")
        ]
        main_b = [
            ("輕量橋式 (Light Hip Bridge)", "2 組 × 8 下 (臀肌喚醒)"),
            ("輕量登階 (Light Step-up)", "2 組 × 5 下／邊 (動作路徑複習)"),
            ("帕羅夫壓 (Pallof press)", "2 組 × 8 下／邊 (核心穩定)"),
            ("髖屈肌伸展 (Hip Flexor Stretch)", "2 組 × 30 秒／邊 (全身放鬆)")
        ]
        short = [
            ("死蟲式 (Dead bug)", "2 組 × 6 下／邊"),
            ("側棒式 (Side Plank)", "2 組 × 20 秒／邊"),
            ("鳥狗式 (Bird dog)", "2 組 × 6 下／邊")
        ]

    doc.add_paragraph(f"四、 肌力訓練動作內容 ({phase_title})", style="Heading 1")
    exercise_table(doc, f"主課 {variant} (週一與週四執行 - {phase_title.split(' ')[0]})", main_a if variant == "A" else main_b)
    exercise_table(doc, "短課模板 (週三執行)", short)
    
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    font(note.add_run("⚠️ 肌力調整規則：若出現關節疼痛、嚴重疲勞或睡眠不足，應果斷將主課減至 2 組或改做輕度活動度拉伸。若有尖銳疼痛請立即停止訓練。"), 10, True, "9C3D10")
    
    # ------------------ SECTION 5 ------------------
    doc.add_paragraph("五、 220 Triathlon 最新鐵人訓練新知與建議", style="Heading 1")
    advice_p = doc.add_paragraph()
    font(advice_p.add_run("結合國際知名鐵人三項網站 "), 11)
    font(advice_p.add_run("220 Triathlon"), 11, True, "1F4D78")
    font(advice_p.add_run(" 的最新訓練新知，為您提供本週與後續訓練的關鍵建議：\n\n"), 11)
    
    # Bullet points for 220Triathlon advice
    p_adv1 = doc.add_paragraph(style="List Bullet")
    p_adv1.paragraph_format.space_after = Pt(4)
    run1_1 = p_adv1.add_run("基礎期的三大核心任務：")
    run1_1.bold = True
    font(run1_1, 11)
    run1_2 = p_adv1.add_run(" 220 Triathlon 的教練強調，基礎期 (Base Phase) 的主要目標是提高「有氧效率」、「代謝脂肪燃燒能力」以及「肌肉防傷彈性」。本週您的 4 次跑步中有 3 次為 Zone 2 (有氧慢跑)，務必將心率裝在區間內，不要超速，以達最佳脂肪代謝效果。")
    font(run1_2, 11)
    
    p_adv2 = doc.add_paragraph(style="List Bullet")
    p_adv2.paragraph_format.space_after = Pt(4)
    run2_1 = p_adv2.add_run("腸胃耐受性訓練 (Gut Training)：")
    run2_1.bold = True
    font(run2_1, 11)
    run2_2 = p_adv2.add_run(" 營養常被稱為鐵人三項的「第五學科」。220 專欄指出，長距離跑步（如本週日 90 分鐘 Zone 2）是練習「腸胃吸收果膠與水分」的最佳時機。不要等到比賽才用，在每次超過 70 分鐘的訓練中，嘗試每 30–45 分鐘攝取 30g 碳水化合物，訓練小腸的吸收能力，避免比賽日腸胃不適。")
    font(run2_2, 11)
    
    p_adv3 = doc.add_paragraph(style="List Bullet")
    p_adv3.paragraph_format.space_after = Pt(4)
    run3_1 = p_adv3.add_run("傾聽身體而非盲從數據 (Listen to Your Body)：")
    run3_1.bold = True
    font(run3_1, 11)
    run3_2 = p_adv3.add_run(" 當前科技產品發達，但頂尖教練提醒，過度關注手錶數據可能導致焦慮或忽視過度訓練。若本週因 4 次跑步與 2 次主肌力課表感到腿部關節緊繃，應以「體感自覺強度 (RPE)」為準，將跑步配速下調，或將肌力課表強度減半，主動恢復。")
    font(run3_2, 11)
    
    p_adv4 = doc.add_paragraph(style="List Bullet")
    p_adv4.paragraph_format.space_after = Pt(4)
    run4_1 = p_adv4.add_run("夏季高溫水分流失補給：")
    run4_1.bold = True
    font(run4_1, 11)
    run4_2 = p_adv4.add_run(" 七月夏日炎熱，汗液流失極快。教練建議在進行 1 小時以上的跑步時，每 15 分鐘固定小口補水，每小時目標補水量為 500–750ml，並搭配電解質鹽錠，以維持血鈉平衡，避免抽筋。")
    font(run4_2, 11)

    # ------------------ SECTION 6 ------------------
    doc.add_paragraph("六、 IM226 挑戰 11 小時內完賽 (Sub-11) 攻略指引", style="Heading 1")
    doc.add_paragraph("為協助您突破個人最佳紀錄 (11:38) 闖入 11 小時內，結合運動科學與知名鐵人網站 220 Triathlon 的 Sub-11 攻略分析如下：")
    
    doc.add_paragraph("(一) 完賽時間目標拆解 (Sub-11 Target Split vs PB)", style="Heading 3")
    
    table_sub11 = doc.add_table(rows=1, cols=4)
    table_sub11.style = "Table Grid"
    headers_sub11 = ["項目", "2024 普悠瑪 (PB)", "Sub-11 目標配速", "目標削減時間"]
    for idx, text in enumerate(headers_sub11):
        table_sub11.rows[0].cells[idx].text = text
        shade(table_sub11.rows[0].cells[idx], "1F4D78")
        for p in table_sub11.rows[0].cells[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                font(run, 10, True, "FFFFFF")
                
    splits_data = [
        ("游泳 (3.8k)", "01:19:36", "01:12:00 (配速 ~1:53/100m)", "- 7 分 36 秒"),
        ("轉換區 T1", "00:04:47", "00:03:30 (動作流暢，省下無謂時間)", "- 1 分 17 秒"),
        ("自行車 (180k)", "05:37:55", "05:25:00 (均速 ~33.2 km/h)", "- 12 分 55 秒"),
        ("轉換區 T2", "00:06:04", "00:03:30 (快速穿鞋，省下無謂時間)", "- 2 分 34 秒"),
        ("跑步 (42.2k)", "04:30:01", "04:11:00 (配速 ~5:56/km)", "- 19 分 01 秒"),
        ("總計 (Total)", "11:38:24", "10:55:00", "- 43 分 24 秒")
    ]
    
    for row_data in splits_data:
        cells = table_sub11.add_row().cells
        for idx, val in enumerate(row_data):
            cells[idx].text = val
            
    table_geometry(table_sub11, [1800, 2000, 3200, 2410])
    for row_idx, row in enumerate(table_sub11.rows):
        if row_idx == 0: continue
        for col_idx, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    font(run, 9.5, False)
                    
    doc.add_paragraph("🎯 重點突破口分析：\n跑步（Marathon）是您的最大突破點！", style="Heading 3")
    p_break = doc.add_paragraph()
    font(p_break.add_run("您在 Puyuma 的自行車成績已達極高水準的 5:37 (均速 32 km/h)，但這也可能造成跑步段掉速至 4:30。若能在自行車項目保持節制，並大幅增強跑步的「關節衝擊耐受性」，將跑步成績拉到 4:10 內，即可輕鬆跨過 11 小時大關。轉換區 (T1 & T2) 同樣有合計約 4 分鐘的免費速度可以壓榨！"))
    
    doc.add_paragraph("(二) 訓練與配速策略", style="Heading 3")
    
    p_strat1 = doc.add_paragraph(style="List Bullet")
    font(p_strat1.add_run("游泳 (Swim) - 「跟車」與流暢度："), 11, True)
    font(p_strat1.add_run(" 220 Triathlon 的公開水域專欄指出，完美的跟游 (Drafting) 能為您省下 18-25% 的阻力與體能。不要孤軍奮戰，在出水前最後 200 米稍微加快打腿頻率，活化腿部血管，為 T1 和騎車做準備。\n"), 11)
    
    p_strat2 = doc.add_paragraph(style="List Bullet")
    font(p_strat2.add_run("自行車 (Bike) - FTP 節制控制："), 11, True)
    font(p_strat2.add_run(" 根據 Best Bike Split 的功率模型，Sub-11 選手的單車強度因子 (IF) 應嚴格控制在 65% - 72% FTP。在上坡、逆風時，切忌暴衝（維持心率在 Zone 2/3 邊界）。保留腿部力量，是保障跑步段不抽筋、維持配速的前提。\n"), 11)
    
    p_strat3 = doc.add_paragraph(style="List Bullet")
    font(p_strat3.add_run("跑步 (Run) - 「鐵人小步跑」防傷："), 11, True)
    font(p_strat3.add_run(" 80/20 Endurance 指出，在 226 最後的全馬中，核心在於「步頻」而非步幅，練習低騰空、低衝擊的 Ironman Shuffle，保持每分鐘 175-180 的步頻，起跑前兩公里刻意壓低配速，防止大腿提早力竭。\n"), 11)
    
    p_strat4 = doc.add_paragraph(style="List Bullet")
    font(p_strat4.add_run("肌力訓練 (Strength) - 姿勢維持力："), 11, True)
    font(p_strat4.add_run(" 220 Triathlon 的肌力專題指出，跑步後半程姿勢崩潰是掉速與受傷主因。本週的單腳硬舉與分腿蹲，旨在訓練您的單腿著地剛性；而側棒式與帕羅夫壓則能加強核心的側向抗扭轉，維持跑步時的骨盆水平。\n"), 11)

    doc.add_paragraph("(三) 關鍵補給與水合策略 (Nutrition & Hydration)", style="Heading 3")
    
    p_nut1 = doc.add_paragraph(style="List Bullet")
    font(p_nut1.add_run("糖分吸收目標 80-90g/hr："), 11, True)
    font(p_nut1.add_run(" 220 Triathlon 指明，要挑戰 Sub-11，補給不可有任何折扣。在自行車階段，每小時必須攝取 80–90 克的碳水化合物（如 3-4 包能量果膠或高碳水運動飲料）。請在週六、週日的長距離練習中反覆練習「腸胃耐受性訓練 (Gut Training)」，讓身體適應高碳水吸收速率。\n"), 11)
    
    p_nut2 = doc.add_paragraph(style="List Bullet")
    font(p_nut2.add_run("多重傳輸糖比例："), 11, True)
    font(p_nut2.add_run(" 建議選用 maltodextrin : fructose (葡萄糖/麥芽糊精 : 果糖) 比例為 1:0.8 或 2:1 的產品，透過不同的腸道通道吸收，將吸收率最大化並降低腸胃發炎機率。\n"), 11)
    
    p_nut3 = doc.add_paragraph(style="List Bullet")
    font(p_nut3.add_run("跑步段避開固體補給："), 11, True)
    font(p_nut3.add_run(" 根據 Roadman Cycling 研究，跑步時血液集中在四肢肌肉，胃部消化能力極低。在路跑段一律使用液體、果膠或軟糖補給，切忌在水站攝取難消化的固體食物，避免造成消化不良或嘔吐。\n"), 11)
    
    p_nut4 = doc.add_paragraph(style="List Bullet")
    font(p_nut4.add_run("水合與鈉離子平衡："), 11, True)
    font(p_nut4.add_run(" 每小時目標攝取 500-750ml 水分，並依據個人汗流率，每升水分中應含有 500-1000mg 的鈉。出汗量大時需加服鹽錠。\n"), 11)

    doc.add_paragraph("(四) 最新運動科學新知 (Sports Science News)", style="Heading 3")
    p_sci1 = doc.add_paragraph(style="List Bullet")
    font(p_sci1.add_run("高碳水吸收新突破："), 11, True)
    font(p_sci1.add_run(" 根據 Precision Fuel & Hydration 最新研究，高階長距離運動員若經過適當的「腸胃訓練」，每小時的碳水化合物吸收量可提升至 100-120g。這能為後半程馬拉松提供源源不絕的肝醣來源，徹底擺脫撞牆期，這是 Sub-11 選手的重要分水嶺。"), 11)

    week_num = monday.isocalendar().week
    
    # Save Streamlined Strength Workout Markdown with Tutorial Links
    main_rows_md = ""
    if variant == "A":
        main_rows_md = """| **分腿蹲 (Split Squat)** | 3 組 × 6–8 下／邊 (強化單腿力量與平衡) | [▶️ 觀看教學 ↗](https://www.youtube.com/results?search_query=分腿蹲+Split+Squat+教學) |
| **羅馬尼亞硬舉 (RDL)** | 3 組 × 6–8 下 (強化臀大肌與大腿後側膕繩肌) | [▶️ 觀看教學 ↗](https://www.youtube.com/results?search_query=羅馬尼亞硬舉+RDL+教學) |
| **單腳硬舉 (Single Leg RDL)** | 2 組 × 6–8 下／邊 (訓練腳踝與髖關節穩定) | [▶️ 觀看教學 ↗](https://www.youtube.com/results?search_query=單腳硬舉+Single+Leg+RDL+教學) |
| **側棒式 (Side Plank)** | 2 組 × 30–45 秒／邊 (強化核心側向抗側彎能力) | [▶️ 觀看教學 ↗](https://www.youtube.com/results?search_query=側棒式+Side+Plank+教學) |
| **死蟲式 (Dead Bug)** | 2 組 × 8–10 下／邊 (核心抗伸展能力) | [▶️ 觀看教學 ↗](https://www.youtube.com/results?search_query=死蟲式+Dead+Bug+教學) |
| **雙腳提踵 (Calf Raises)** | 3 組 × 12–15 下 (加強小腿與跟腱強度) | [▶️ 觀看教學 ↗](https://www.youtube.com/results?search_query=提踵+Calf+Raises+教學) |"""
    else:
        main_rows_md = """| **臀推或橋式 (Hip Thrust)** | 3 組 × 8–10 下 (活化並強化臀大肌) | [▶️ 觀看教學 ↗](https://www.youtube.com/results?search_query=臀推+Hip+Thrust+教學) |
| **登階 (Step-up)** | 3 組 × 6–8 下／边 (模仿跑步推蹬與單腳穩定) | [▶️ 觀看教學 ↗](https://www.youtube.com/results?search_query=登階+Step-up+教學) |
| **帕羅夫壓 (Pallof Press)** | 2 組 × 10–12 下／邊 (核心抗旋轉能力) | [▶️ 觀看教學 ↗](https://www.youtube.com/results?search_query=帕羅夫壓+Pallof+Press+教學) |
| **怪獸走 (Monster Walk)** | 2 組 × 10–15 步／邊 (激活臀中肌，防止膝外翻) | [▶️ 觀看教學 ↗](https://www.youtube.com/results?search_query=怪獸走+Monster+Walk+教學) |
| **離心提踵 (Eccentric Calf)** | 3 組 × 8–12 下 (預防跟腱炎，加強推蹬剛性) | [▶️ 觀看教學 ↗](https://www.youtube.com/results?search_query=離心提踵+Eccentric+Calf+教學) |
| **髖屈肌伸展 (Hip Flexor Stretch)** | 2 組 × 30 秒／邊 (舒緩久坐與騎車的緊繃) | [▶️ 觀看教學 ↗](https://www.youtube.com/results?search_query=髖屈肌伸展+Hip+Flexor+Stretch+教學) |"""

    md_content = f"""# 🏋️ 2026-W{week_num:02d} 第 {week_num} 週肌力訓練課表詳細內容

### 📌 本週肌力訓練安排
- **週一 & 週四（主課 {variant}）：** 專注下肢多關節推拉與髖關節主導動作（30–40 分鐘），保留 2–3 下餘裕。
- **週三（短課模板）：** 游泳後或晚間進行核心抗旋轉、抗側彎與足踝推蹬剛性訓練（15–20 分鐘）。
- **週日（恢復與拉伸）：** 全身滾筒放鬆與胸椎/髖關節活動度拉伸。

---

### 🏋️ 主課 {variant}（週一與週四執行，約 30–40 分鐘）

| 動作名稱 | 訓練目標與建議量 | 🎥 動作教學影片 |
| :--- | :--- | :---: |
{main_rows_md}

---

### ⚡ 短課模板（週三執行，約 15–20 分鐘）

| 動作名稱 | 訓練目標與建議量 | 🎥 動作教學影片 |
| :--- | :--- | :---: |
| **死蟲式 (Dead Bug)** | 2 組 × 8–10 下／邊 (核心抗伸展能力) | [▶️ 觀看教學 ↗](https://www.youtube.com/results?search_query=死蟲式+Dead+Bug+教學) |
| **側棒式 (Side Plank)** | 2 組 × 30 秒／邊 (強化核心抗側彎) | [▶️ 觀看教學 ↗](https://www.youtube.com/results?search_query=側棒式+Side+Plank+教學) |
| **鳥狗式 (Bird Dog)** | 2 組 × 8 下／邊 (對側核心控制力) | [▶️ 觀看教學 ↗](https://www.youtube.com/results?search_query=鳥狗式+Bird+Dog+教學) |
| **彈力帶側走 (Band Walk)** | 2 組 × 10–15 步／邊 (臀中肌激活) | [▶️ 觀看教學 ↗](https://www.youtube.com/results?search_query=彈力帶側走+Band+Walk+教學) |
| **單腳提踵 (Single Calf Raise)** | 2 組 × 12–15 下／邊 (足踝推蹬剛性) | [▶️ 觀看教學 ↗](https://www.youtube.com/results?search_query=單腳提踵+Single+Calf+Raise+教學) |

---

> ⚠️ **肌力調整與安全規則：**
> 1. 若出現關節疼痛、嚴重疲勞或睡眠不足，應果斷將主課減至 2 組或改做輕度活動度拉伸。
> 2. 若出現尖銳疼痛，請立即停止該項訓練。
> 3. 離心動作（如提踵）請維持 3 秒緩慢下降，勿快速摔落下彈。

"""
    md_file1 = OUT_DIR / f"{monday:%Y}-W{week_num:02d}_第{week_num:02d}週肌力訓練計畫.md"
    md_file2 = OUT_DIR / f"{monday:%Y}-W{week_num:02d}_當週肌力訓練計畫.md"
    md_file1.write_text(md_content, encoding="utf-8")
    md_file2.write_text(md_content, encoding="utf-8")

    output = OUT_DIR / f"{monday:%Y}-W{week_num:02d}_第{week_num:02d}週肌力訓練計畫.docx"
    try:
        doc.save(output)
        print(f"Successfully generated merged report and strength document at: {output}")
    except PermissionError:
        import time
        ts = int(time.time())
        fallback_output = OUT_DIR / f"{monday:%Y}-W{week_num:02d}_第{week_num:02d}週肌力訓練計畫_{ts}.docx"
        doc.save(fallback_output)
        print(f"File locked, successfully saved with timestamp to: {fallback_output}")
        return fallback_output
        
    return output


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--week", help="該週任一天，格式 YYYY-MM-DD")
    args = parser.parse_args()
    anchor = date.fromisoformat(args.week) if args.week else date.today()
    monday = anchor - timedelta(days=anchor.weekday())
    
    # Run sync automatically first to update local database with the latest events
    try:
        from sync_calendar import sync
        print("Syncing calendar to cache first...")
        sync()
    except Exception as e:
        print("Sync skipped or failed:", e)
        
    build(monday)


if __name__ == "__main__":
    main()
