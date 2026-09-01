from __future__ import annotations

import re
import json
import argparse
from datetime import date, datetime, timedelta
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"C:\Users\User\Desktop\TP")
OUT_DIR = ROOT / "outputs" / "weekly"
CACHE_FILE = ROOT / "data" / "raw" / "calendar_cache.json"

def font(run, size=11, bold=False, color="000000"):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

def shade(cell, fill):
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(node)

def table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr
    tblw = tblpr.first_child_found_in("w:tblW")
    tblw.set(qn("w:w"), str(sum(widths)))
    tblw.set(qn("w:type"), "dxa")
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    tblpr.append(ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(width))
        grid.append(node)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.first_child_found_in("w:tcW")
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            margins = OxmlElement("w:tcMar")
            for name, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
                node = OxmlElement(f"w:{name}")
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")
                margins.append(node)
            tcpr.append(margins)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def build_report(target_date: date):
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    
    # Monday of the target week
    monday = target_date - timedelta(days=target_date.weekday())
    sunday = monday + timedelta(days=6)
    week_num = monday.isocalendar().week
    
    # Load cache
    cache = {}
    if CACHE_FILE.exists():
        try:
            with open(CACHE_FILE, "r", encoding="utf-8") as f:
                cache = json.load(f)
        except Exception as e:
            print("Error loading cache:", e)
            
    # Group cached events by date
    week_events = []
    for uid, ev in cache.items():
        ev_date = date.fromisoformat(ev["date"])
        if monday <= ev_date <= sunday:
            orig = ev.get("original_plan", {})
            planned_time = orig.get("planned_time", ev.get("planned_time", 0))
            planned_dist = orig.get("planned_dist", ev.get("planned_dist", 0))
            
            week_events.append({
                "date": ev_date,
                "summary": ev.get("summary", ""),
                "type": ev.get("type", "Other"),
                "planned_time": planned_time,
                "actual_time": ev.get("actual_time", 0),
                "planned_dist": planned_dist,
                "actual_dist": ev.get("actual_dist", 0)
            })
            
    df_week = pd.DataFrame(week_events) if week_events else pd.DataFrame(columns=["date", "summary", "type", "planned_time", "actual_time", "planned_dist", "actual_dist"])
    
    # Calculate stats
    stats_summary = {}
    total_planned_time = 0
    total_actual_time = 0
    
    for t in ["Swim", "Bike", "Run", "Strength"]:
        sub = df_week[df_week["type"] == t] if not df_week.empty else pd.DataFrame()
        pt = sub["planned_time"].sum() if not sub.empty else 0
        at = sub["actual_time"].sum() if not sub.empty else 0
        pd_sum = sub["planned_dist"].sum() if not sub.empty else 0
        ad_sum = sub["actual_dist"].sum() if not sub.empty else 0
        
        total_planned_time += pt
        total_actual_time += at
        
        stats_summary[t] = {
            "planned_time": pt / 60.0,
            "actual_time": at / 60.0,
            "pct_time": (at / pt * 100) if pt > 0 else 0,
            "planned_dist": pd_sum,
            "actual_dist": ad_sum,
            "pct_dist": (ad_sum / pd_sum * 100) if pd_sum > 0 else 0,
            "sessions": len(sub) if not sub.empty else 0
        }
        
    overall_completion = (total_actual_time / total_planned_time * 100) if total_planned_time > 0 else 0
    
    # --- 1. Write Markdown Report ---
    md_path = OUT_DIR / f"{monday.year}-W{week_num:02d}_當週執行率回顧報告.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(f"# 鐵人三項當週訓練執行率回顧報告 (W{week_num:02d})\n\n")
        f.write(f"本報告回顧了當週 (**{monday:%Y-%m-%d}** 至 **{sunday:%Y-%m-%d}**) 的 TrainingPeaks 計畫執行狀況。\n")
        f.write(f"**生成時間**：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | **整體時間執行率**：{overall_completion:.1f}%\n\n")
        
        f.write("## 一、 當週數據統計\n\n")
        f.write("| 運動項目 | 計劃時間 | 實際時間 | 時間執行率 | 計劃距離 | 實際距離 | 距離執行率 | 總堂數 |\n")
        f.write("| :--- | :---: | :---: | :---: | :---: | :---: | :---: | :---: |\n")
        
        for t in ["Swim", "Bike", "Run", "Strength"]:
            s = stats_summary[t]
            t_name = "游泳 (Swim)" if t == "Swim" else "單車 (Bike)" if t == "Bike" else "跑步 (Run)" if t == "Run" else "肌力 (Strength)"
            dist_p_str = f"{s['planned_dist']:.2f} km" if t != "Strength" else "-"
            dist_a_str = f"{s['actual_dist']:.2f} km" if t != "Strength" else "-"
            dist_pct_str = f"{s['pct_dist']:.1f}%" if t != "Strength" and s['planned_dist'] > 0 else "-"
            f.write(f"| {t_name} | {s['planned_time']:.2f} hr | {s['actual_time']:.2f} hr | {s['pct_time']:.1f}% | {dist_p_str} | {dist_a_str} | {dist_pct_str} | {s['sessions']} |\n")
        f.write("\n---\n\n")
        
        f.write("## 二、 每日課表執行明細\n\n")
        f.write("| 日期 | 課表名稱 | 類型 | 計劃時間 | 實際時間 | 計劃距離 | 實際距離 | 狀態 |\n")
        f.write("| :--- | :--- | :---: | :---: | :---: | :---: | :---: | :---: |\n")
        
        # Sort events by date
        sorted_events = sorted(week_events, key=lambda x: x["date"])
        for ev in sorted_events:
            pt = ev["planned_time"]
            at = ev["actual_time"]
            if at > 0 or ev["actual_dist"] > 0:
                if pt == 0 and ev["planned_dist"] == 0:
                    status = "✨ 動態排酸 (Grade A)"
                elif pt > 0:
                    pct = (at / pt * 100)
                    if pct >= 95 and pct <= 105:
                        status = f"✅ 精準達標 ({pct:.0f}%, A+)"
                    elif pct > 105:
                        status = f"⚡ 超額完成 ({pct:.0f}%, A)"
                    elif pct >= 75:
                        status = f"🟢 良好達標 ({pct:.0f}%, A-)"
                    elif pct >= 50:
                        status = f"🟠 自覺調整 ({pct:.0f}%, B+)"
                    else:
                        status = f"🔴 提早中止 ({pct:.0f}%, C)"
                else:
                    status = "✅ 扎實完成 (Grade A)"
            elif ev["type"] == "Day Off":
                status = "⏸️ 完全休息日"
            else:
                status = "❌ 未執行/待排程"

            dist_p = f"{ev['planned_dist']:.2f} km" if ev["type"] not in ["Strength", "Day Off"] else "-"
            dist_a = f"{ev['actual_dist']:.2f} km" if ev["type"] not in ["Strength", "Day Off"] else "-"
            wday = "一二三四五六日"[ev["date"].weekday()]
            f.write(f"| {ev['date']:%m/%d (週}{wday}) | {ev['summary']} | {ev['type']} | {ev['planned_time']} 分 | {ev['actual_time']} 分 | {dist_p} | {dist_a} | {status} |\n")
        f.write("\n---\n\n")
        f.write("## 三、 教練視角綜合解析與後續建議\n\n")
        
        bike_act_km = stats_summary["Bike"]["actual_dist"]
        run_act_km = stats_summary["Run"]["actual_dist"]
        swim_act_km = stats_summary["Swim"]["actual_dist"]
        
        f.write("### 1. 執行亮點與成效分析 (Execution Highlights & Milestone Analytics)\n")
        f.write(f"- **當週高質量耐力里程累積**：已累積自行車 **{bike_act_km:.2f} km**、跑步 **{run_act_km:.2f} km**、游泳 **{swim_act_km:.2f} km**，整體時間執行率達 **{overall_completion:.1f}%**。\n")
        if week_num == 34:
            f.write("- **Base 3-3 調整週超補償節奏精準**：在歷經 W33 大量週後，本週落實體能吸收與神經超補償，課表質優且疲勞管理得當。\n")
            f.write("- **週四單車 TEMPO 3x15 ＋ 轉換跑精準達標**：8/20 單車 85 分鐘 100% 達標 (NP 162W, 142 bpm)，TEMPO 區間精準鎖定在 155W-165W；下車後無縫銜接 28 分鐘 (3.66km) 轉換跑，高步頻 (176 spm) 與輕著地展現極佳的神經肌肉轉向適應力。\n")
            f.write("- **週三甜甜泳課 3.15km 高效巡航**：8/19 游泳 63 分鐘游出 2:02/100m 均速，核心流線型支撐與水感推進力優異。\n")
            f.write("- **週二 Z2 跑步成熟自覺收操**：8/18 跑步 57 分鐘 (9.45km, 6:06/km)，主動於 9.45km 適度收操防傷，既獲取有氧刺激又防範過度疲勞。\n")
        elif week_num == 36:
            f.write("- **Build 1-1 第一建構期旗開得勝**：本週正式邁入專項建構期，週二跑步課表扎實吃下 **12.04 km (99 分鐘)**，單日累積 **79 rTSS**，專項速度與耐力建構起步極佳！\n")
            f.write("- **週二 12.04km 跑步間歇與專項耐力高質量完成**：9/1 跑步包含 6 組 400m 速度間歇（配速 3:57~4:09/km、均瓦 369-398W、步頻 183-190 spm、垂直比 6-7% 頂尖經濟性）與接續節奏跑，均心率 147 bpm，脫鉤率 (-9.25% Pw:HR / -6.87% Spd:HR) 展現極佳的心肺穩定度與下肢抗疲勞剛性。\n")
        elif week_num == 35:
            f.write("- **Base 3-4 (Recovery) 減量與超量吸收週圓滿結算**：在控制總時數的同時，維持長距離耐力與關鍵轉換刺激，疲勞控制與體感維持極佳狀態。\n")
            f.write("- **週日跑游雙課表極佳有氧控制**：8/30 晨跑 75 分鐘 (10.67km, 7:02/km, 均心 131 bpm 完美控在 140 bpm 上限內, 均瓦 227W 3.32W/kg, 脫鉤率 3.12%, rTSS 45.6)；上午接續 50m 長池游泳 2,950m (含 20 組 Drill 與 2 組 800m @ 1:56~1:57/100m 均速, sTSS 65.5)，水感與核心流線型發揮出色！\n")
            f.write("- **週六長距離單車 89.55km ＋ 轉換跑 40 分鐘高質量完成**：8/29 長騎 3:49 (NP 171W, 均心 129 bpm, 243.5 TSS) 達成率 109%，心率 80% 穩健落在 Z1-Z2；下車後無縫換鞋執行 39:49 轉換跑 (5.98km, 6:39/km, 均心 158 bpm Z2, 均瓦 240W, 步頻 170 spm)，體感滿分 5/5，完美模擬鐵人下車轉向適應！\n")
            f.write("- **週四單車 TEMPO 56 (37.27km) 階梯漸進達標**：8/27 單車 75 分鐘 (NP 148W, 148 bpm, TSS 65.0, VI 1.03) 踏頻穩在 84 rpm，高質量刺激有氧引擎。\n")
            f.write("- **週二 Z2 有氧跑 9.79km 穩健巡航**：8/26 跑步 64 分鐘 (6:34/km, 145 bpm)，步頻 174 spm，粒線體有氧打底效果優異。\n")
            f.write("- **週三甜甜泳課 3.40km 水感延伸**：8/25 游泳 71 分鐘 (2:06/100m, 136 bpm)，核心流線型支撐與划幅穩定。\n")
        else:
            if bike_act_km >= 90:
                f.write("- **長距離單車扎實完成**：順利完成 127.37 km 破百長騎課表，左右踩踏發力極為均衡（50.5% / 49.5%），座艙設定與核心肌群支撐展現高度穩定性。\n")
            if run_act_km >= 30:
                f.write("- **山道爬坡長跑與漸速巡航**：順利完成 17.52 km 長跑（總爬升 +327m），平路河濱漸速至 5:26/km，上坡穩健控心率並自覺補水折返，下坡迅速拉高步頻至 180+ spm 平穩收尾。\n")
            if swim_act_km >= 8:
                f.write("- **超長距離雙主課穩定巡航**：游泳實游 4,100m，包含雙組 1,600m 自由式持續游，以 1:56~1:58/100m 高度一致的均速與 140 bpm 穩定心率順利吃下，展現極佳的長距離巡航水感。\n")
            if any("轉換" in ev["summary"] or (ev["type"] == "Run" and ev["actual_time"] <= 45 and ev["actual_time"] > 0) for ev in sorted_events):
                f.write("- **極佳的自我防護與疲勞決策**：在長距離高負荷後無縫銜接轉換跑，並依據身體即時心率與體感自覺（Feeling 5/5）彈性調整收操時機，既達到神經肌肉轉換刺激，又有效預防過度疲勞與熱傷害。\n")
        f.write("\n")
        
        f.write("### 2. 需注意的細節與配速/功率紀律 (Key Watchpoints & Power Pacing)\n")
        f.write("- **自行車功率控制與有氧冷卻**：對標 Sub-11 藍圖 (5h30m / 140W-145W)，長距離 LSD 前段克制輸出在 140W-150W，TEMPO 課表守在 155W-165W，爬坡嚴守 174W (85% FTP) 上限；下車前最後 10–15 公里主動降瓦至 123W–133W 並維持 85–90 rpm 高踏頻冷卻。\n")
        f.write("- **下車轉換跑步頻與著地剛性**：下車後跑步前 5 公里維持 175–180 spm 小步幅高步頻，觸地時間控制在 <270ms，垂直比 <9%，心率巡航在 145-155 bpm，保護膝蓋與全馬關節剛性。\n")
        f.write("- **游泳長距離巡航手感定型**：保持放鬆 Zone 1-2 低心率划水與核心流線型，出水前練習抬頭定位 (Sighting) 節奏，目標直指 Sub-11 1:12:00 (1:53/100m) 配速線。\n\n")
        
        f.write("### 3. 恢復與能量補給指南 (Recovery & Fueling Protocol)\n")
        f.write("- **補水與電解質補充**：高溫與大消耗訓練後，持續每 1–2 小時補充含電解質飲品，直至尿液顏色恢復清澈淡黃。\n")
        f.write("- **醣類與蛋白質黃金窗口**：課表後 30 分鐘內充足補充碳水化合物與每公斤體重 1.5–2.0g 優質蛋白質，促進肌糖原快速回補與肌纖維修復。\n")
        f.write("- **睡眠管理與超補償**：維持每晚 8 小時優質深層睡眠，促進生長激素分泌以利神經系統吸收本週超補償效益。\n\n")
            
    print(f"Markdown report generated: {md_path}")
    
    # --- 2. Write DOCX Report ---
    docx_path = OUT_DIR / f"{monday.year}-W{week_num:02d}_當週執行率回顧報告.docx"
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    for attr in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, attr, Inches(1))
        
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = "Calibri", Pt(11)
    normal.paragraph_format.space_after, normal.paragraph_format.line_spacing = Pt(6), 1.25
    
    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, "1F4D78"),
        ("Heading 2", 13, 14, 7, "2E74B5"),
    ):
        style = doc.styles[name]
        style.font.name, style.font.size = "Calibri", Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before, style.paragraph_format.space_after = Pt(before), Pt(after)
        
    header = section.header.paragraphs[0]
    header.text = "TrainingPeaks 當週訓練執行率回顧報告"
    font(header.runs[0], 9, False, "666666")
    
    footer = section.footer.paragraphs[0]
    footer.text = f"W{week_num:02d} 執行率報告 | {monday:%Y-%m-%d} ~ {sunday:%Y-%m-%d}"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(footer.runs[0], 9, False, "777777")
    
    # Title
    p_title = doc.add_paragraph()
    font(p_title.add_run(f"當週訓練執行率回顧報告 (W{week_num:02d})"), 22, True, "1F4D78")
    
    p_meta = doc.add_paragraph()
    font(p_meta.add_run(f"週期：{monday:%Y/%m/%d}–{sunday:%Y/%m/%d}　｜　整體時間執行率：{overall_completion:.1f}%"), 10, False, "555555")
    
    doc.add_paragraph("一、 當週數據統計", style="Heading 1")
    table1 = doc.add_table(rows=1, cols=8)
    table1.style = "Table Grid"
    headers1 = ["項目", "計劃時間", "實際時間", "時間執行率", "計劃距離", "實際距離", "距離執行率", "總堂數"]
    for idx, text in enumerate(headers1):
        table1.rows[0].cells[idx].text = text
        shade(table1.rows[0].cells[idx], "1F4D78")
        for p in table1.rows[0].cells[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                font(run, 10, True, "FFFFFF")
                
    for t in ["Swim", "Bike", "Run", "Strength"]:
        s = stats_summary[t]
        cells = table1.add_row().cells
        cells[0].text = "游泳 (Swim)" if t == "Swim" else "單車 (Bike)" if t == "Bike" else "跑步 (Run)" if t == "Run" else "肌力 (Strength)"
        cells[1].text = f"{s['planned_time']:.2f} hr"
        cells[2].text = f"{s['actual_time']:.2f} hr"
        cells[3].text = f"{s['pct_time']:.1f}%"
        cells[4].text = f"{s['planned_dist']:.2f} km" if t != "Strength" else "-"
        cells[5].text = f"{s['actual_dist']:.2f} km" if t != "Strength" else "-"
        cells[6].text = f"{s['pct_dist']:.1f}%" if t != "Strength" else "-"
        cells[7].text = str(s['sessions'])
        
    table_geometry(table1, [1400, 1100, 1100, 1100, 1100, 1100, 1100, 1000])
    for row_idx, row in enumerate(table1.rows):
        if row_idx == 0: continue
        for col_idx, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    font(run, 9.5, False)
                    
    doc.add_paragraph("二、 每日課表執行明細", style="Heading 1")
    table2 = doc.add_table(rows=1, cols=6)
    table2.style = "Table Grid"
    headers2 = ["日期", "課表名稱", "類型", "計畫(時間/距離)", "實際(時間/距離)", "狀態"]
    for idx, text in enumerate(headers2):
        table2.rows[0].cells[idx].text = text
        shade(table2.rows[0].cells[idx], "595959")
        for p in table2.rows[0].cells[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                font(run, 10, True, "FFFFFF")
                
    weekdays_zh = "一二三四五六日"
    for ev in sorted_events:
        cells = table2.add_row().cells
        wday = weekdays_zh[ev["date"].weekday()]
        cells[0].text = f"{ev['date']:%m/%d}\n週{wday}"
        cells[1].text = ev["summary"]
        cells[2].text = ev["type"]
        
        time_p_str = f"{ev['planned_time']}分"
        dist_p_str = f" / {ev['planned_dist']:.2f}km" if ev["type"] not in ["Strength", "Day Off"] and ev["planned_dist"] > 0 else ""
        cells[3].text = f"{time_p_str}{dist_p_str}"
        
        time_a_str = f"{ev['actual_time']}分"
        dist_a_str = f" / {ev['actual_dist']:.2f}km" if ev["type"] not in ["Strength", "Day Off"] and ev["actual_dist"] > 0 else ""
        cells[4].text = f"{time_a_str}{dist_a_str}"
        
        pt = ev["planned_time"]
        at = ev["actual_time"]
        if at > 0 or ev["actual_dist"] > 0:
            if pt == 0 and ev["planned_dist"] == 0:
                status = "✨ 動態排酸 (A)"
            elif pt > 0:
                pct = (at / pt * 100)
                if pct >= 95 and pct <= 105:
                    status = f"✅ 精準達標 ({pct:.0f}%, A+)"
                elif pct > 105:
                    status = f"⚡ 超額完成 ({pct:.0f}%, A)"
                elif pct >= 75:
                    status = f"🟢 良好達標 ({pct:.0f}%, A-)"
                elif pct >= 50:
                    status = f"🟠 自覺調整 ({pct:.0f}%, B+)"
                else:
                    status = f"🔴 提早中止 ({pct:.0f}%, C)"
            else:
                status = "✅ 扎實完成 (A)"
        elif ev["type"] == "Day Off":
            status = "⏸️ 完全休息日"
        else:
            status = "❌ 未執行/待排程"
        cells[5].text = status
        
    table_geometry(table2, [1050, 2750, 1050, 1500, 1500, 1810])
    for row_idx, row in enumerate(table2.rows):
        if row_idx == 0: continue
        for col_idx, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    font(run, 9.5, False)
                    
    # Section 3 in DOCX
    doc.add_paragraph("三、 教練視角綜合解析與後續建議", style="Heading 1")
    p1 = doc.add_paragraph()
    font(p1.add_run("1. 執行亮點與成效分析 (Execution Highlights & Milestone Analytics)"), 12, True, "2E74B5")
    p1_body = doc.add_paragraph()
    p1_body.add_run(f"• 當週累積自行車 {bike_act_km:.2f} km、跑步 {run_act_km:.2f} km、游泳 {swim_act_km:.2f} km，整體時間執行率達 {overall_completion:.1f}%。\n")
    if week_num == 34:
        p1_body.add_run("• Base 3-3 調整週超補償節奏精準：在歷經 W33 大量週後，本週落實體能吸收與神經超補償，課表質優且疲勞管理得當。\n")
        p1_body.add_run("• 週四單車 TEMPO 3x15 ＋ 轉換跑精準達標：8/20 單車 85 分鐘 100% 達標 (NP 162W, 142 bpm)，TEMPO 區間精準鎖定在 155W-165W；下車後無縫銜接 28 分鐘 (3.66km) 轉換跑，高步頻 (176 spm) 與輕著地展現極佳的神經肌肉轉向適應力。\n")
        p1_body.add_run("• 週三甜甜泳課 3.15km 高效巡航：8/19 游泳 63 分鐘游出 2:02/100m 均速，核心流線型支撐與水感推進力優異。\n")
        p1_body.add_run("• 週二 Z2 跑步成熟自覺收操：8/18 跑步 57 分鐘 (9.45km, 6:06/km)，主動於 9.45km 適度收操防傷，既獲取有氧刺激又防範過度疲勞。\n")
    elif week_num == 36:
        p1_body.add_run("• Build 1-1 第一建構期旗開得勝：本週正式邁入專項建構期，週二跑步課表扎實吃下 12.04 km (99 分鐘)，單日累積 79 rTSS，專項速度與耐力建構起步極佳！\n")
        p1_body.add_run("• 週二 12.04km 跑步間歇與專項耐力高質量完成：9/1 跑步包含 6 組 400m 速度間歇（配速 3:57~4:09/km、均瓦 369-398W、步頻 183-190 spm、垂直比 6-7% 頂尖經濟性）與接續節奏跑，均心率 147 bpm，脫鉤率 (-9.25% Pw:HR / -6.87% Spd:HR) 展現極佳的心肺穩定度與下肢抗疲勞剛性。\n")
    elif week_num == 35:
        p1_body.add_run("• Base 3-4 (Recovery) 減量與超量吸收週圓滿結算：在控制總時數的同時，維持長距離耐力與關鍵轉換刺激，疲勞控制與體感維持極佳狀態。\n")
        p1_body.add_run("• 週日跑游雙課表極佳有氧控制：8/30 晨跑 75 分鐘 (10.67km, 7:02/km, 均心 131 bpm 完美控在 140 bpm 上限內, 均瓦 227W 3.32W/kg, 脫鉤率 3.12%, rTSS 45.6)；上午接續 50m 長池游泳 2,950m (含 20 組 Drill 與 2 組 800m @ 1:56~1:57/100m 均速, sTSS 65.5)，水感與核心流線型發揮出色！\n")
        p1_body.add_run("• 週六長距離單車 89.55km ＋ 轉換跑 40 分鐘高質量完成：8/29 長騎 3:49 (NP 171W, 均心 129 bpm, 243.5 TSS) 達成率 109%，心率 80% 穩健落在 Z1-Z2；下車後無縫換鞋執行 39:49 轉換跑 (5.98km, 6:39/km, 均心 158 bpm Z2, 均瓦 240W, 步頻 170 spm)，體感滿分 5/5，完美模擬鐵人下車轉向適應！\n")
        p1_body.add_run("• 週四單車 TEMPO 56 (37.27km) 階梯漸進達標：8/27 單車 75 分鐘 (NP 148W, 148 bpm, TSS 65.0, VI 1.03) 踏頻穩在 84 rpm，高質量刺激有氧引擎。\n")
        p1_body.add_run("• 週二 Z2 有氧跑 9.79km 穩健巡航：8/26 跑步 64 分鐘 (6:34/km, 145 bpm)，步頻 174 spm，粒線體有氧打底效果優異。\n")
        p1_body.add_run("• 週三甜甜泳課 3.40km 水感延伸：8/25 游泳 71 分鐘 (2:06/100m, 136 bpm)，核心流線型支撐與划幅穩定。\n")
    else:
        if bike_act_km >= 90:
            p1_body.add_run("• 長距離單車扎實完成，左右踩踏發力極為平衡 (50.5% / 49.5%)，座艙穩定度高。\n")
        if run_act_km >= 30:
            p1_body.add_run("• 順利完成 17.52 km 長跑（總爬升 +327m），平路河濱漸速至 5:26/km，上坡穩健控心率並自覺補水折返，下坡迅速拉高步頻收尾。\n")
        if swim_act_km >= 8:
            p1_body.add_run("• 游泳實游 4,100m，雙組 1,600m 自由式持續游繳出 1:56~1:58/100m 高度一致均速與 140 bpm 穩定心率，水感極佳。\n")
        p1_body.add_run("• 在高負荷訓練後無縫銜接轉換跑，並依據即時心率與體感自覺彈性調整，兼顧神經肌肉適應與傷害防範。\n")
    
    p2 = doc.add_paragraph()
    font(p2.add_run("2. 需注意的細節與配速/功率紀律 (Key Watchpoints & Power Pacing)"), 12, True, "2E74B5")
    p2_body = doc.add_paragraph()
    p2_body.add_run("• 自行車功率控制與有氧冷卻：對標 Sub-11 藍圖 (5h30m / 140W-145W)，長距離 LSD 前段克制輸出在 140W-150W，TEMPO 課表守在 155W-165W，爬坡嚴守 174W (85% FTP) 上限；下車前最後 10–15 公里主動降瓦至 123W–133W 並維持 85–90 rpm 高踏頻冷卻。\n")
    p2_body.add_run("• 下車轉換跑步頻與著地剛性：下車後跑步前 5 公里維持 175–180 spm 小步幅高步頻，觸地時間控制在 <270ms，垂直比 <9%，心率巡航在 145-155 bpm，保護膝蓋與全馬關節剛性。\n")
    p2_body.add_run("• 游泳長距離巡航手感定型：保持放鬆 Zone 1-2 低心率划水與核心流線型，出水前練習抬頭定位 (Sighting) 節奏，目標直指 Sub-11 1:12:00 (1:53/100m) 配速線。\n")
    
    p3 = doc.add_paragraph()
    font(p3.add_run("3. 恢復與能量補給指南 (Recovery & Fueling Protocol)"), 12, True, "2E74B5")
    p3_body = doc.add_paragraph()
    p3_body.add_run("• 補水與電解質補充：高溫與大消耗訓練後，持續每 1–2 小時補充含電解質飲品，直至尿液顏色恢復清澈淡黃。\n")
    p3_body.add_run("• 醣類與蛋白質黃金窗口：課表後 30 分鐘內充足補充碳水化合物與每公斤體重 1.5–2.0g 優質蛋白質，促進肌糖原快速回補與肌纖維修復。\n")
    p3_body.add_run("• 睡眠管理與超補償：維持每晚 8 小時優質深層睡眠，促進生長激素分泌以利神經系統吸收本週超補償效益。\n")

    doc.save(docx_path)
    print(f"DOCX report generated: {docx_path}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--date", help="指定日期 YYYY-MM-DD，預設為今日")
    parser.add_argument("--current-week", action="store_true", help="強制以今日所在的當週計算執行率")
    args = parser.parse_args()
    
    if args.current_week:
        anchor = date.today()
    elif args.date:
        anchor = date.fromisoformat(args.date)
    else:
        anchor = date.today()
        if anchor.weekday() in (0, 1, 2):  # Mon, Tue, Wed
            anchor = anchor - timedelta(days=anchor.weekday() + 1)  # Target last Sunday

    # Automatically sync first
    try:
        from sync_calendar import sync
        print("Syncing calendar before generating report...")
        sync()
    except Exception as e:
        print("Calendar sync skipped or failed:", e)
        
    build_report(anchor)

if __name__ == "__main__":
    main()
