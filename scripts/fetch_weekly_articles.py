import urllib.request
import xml.etree.ElementTree as ET
from datetime import datetime, timedelta, timezone
import email.utils
from pathlib import Path
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Feeds Configuration
FEEDS = {
    "Slowtwitch": "https://www.slowtwitch.com/feed/",
    "Triathlete Magazine": "https://www.triathlete.com/feed/",
    "220 Triathlon": "https://www.220triathlon.com/feed/",
    "TrainingPeaks Blog": "https://www.trainingpeaks.com/feed/",
    "Joe Friel Blog": "https://joefrieltraining.com/feed/",
}

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.5",
}

ROOT_DIR = Path(r"C:\Users\User\Desktop\TP")
OUTPUT_DIR = ROOT_DIR / "outputs" / "weekly"

def parse_pub_date(pub_date_str):
    try:
        dt = email.utils.parsedate_to_datetime(pub_date_str)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt
    except Exception:
        return None

def clean_html(raw_html):
    if not raw_html:
        return ""
    cleanr = re.compile('<.*?>|&([a-z0-9]+|#[0-9]+|#x[0-9a-f]+);')
    cleantext = re.sub(cleanr, '', raw_html)
    return cleantext.strip()

def categorize_article(title, description, categories):
    content_text = f"{title} {description} {' '.join(categories)}".lower()
    
    # 游泳 (Swimming)
    swim_keywords = ["swim", "pool", "wetsuit", "water", "stroke", "goggle", "open water", "robe", "orca", "游泳", "防寒衣", "跟游", "蛙鏡"]
    if any(kw in content_text for kw in swim_keywords):
        return "游泳 (Swimming)"
        
    # 騎車 (Cycling)
    bike_keywords = ["bike", "cycle", "ftp", "cadence", "pedal", "wahoo", "garmin", "trainer", "gear", "aero", "fit", "cda", "speedmax", "canyon", "pinarello", "單車", "自行車", "功率", "風阻", "騎車"]
    if any(kw in content_text for kw in bike_keywords):
        return "騎車 (Cycling)"
        
    # 跑步 (Running)
    run_keywords = ["run", "shoe", "marathon", "pace", "cadence", "spm", "saucony", "endorphin", "running", "跑步", "跑鞋", "全馬", "步頻"]
    if any(kw in content_text for kw in run_keywords):
        return "跑步 (Running)"
        
    # 補給與恢復 (Fueling & Recovery)
    return "補給與恢復 (Fueling & Recovery)"

def fetch_recent_articles(days_back=7):
    cutoff_date = datetime.now(timezone.utc) - timedelta(days=days_back)
    all_articles = []
    
    for source_name, url in FEEDS.items():
        print(f"Fetching articles from {source_name}...")
        try:
            req = urllib.request.Request(url, headers=HEADERS)
            with urllib.request.urlopen(req, timeout=20) as resp:
                data = resp.read()
                root = ET.fromstring(data)
                
                items = root.findall(".//item")
                source_count = 0
                for item in items:
                    title_node = item.find("title")
                    link_node = item.find("link")
                    pub_date_node = item.find("pubDate")
                    desc_node = item.find("description")
                    
                    title = title_node.text.strip() if title_node is not None else "Untitled"
                    link = link_node.text.strip() if link_node is not None else ""
                    pub_date_str = pub_date_node.text.strip() if pub_date_node is not None else ""
                    description = clean_html(desc_node.text) if desc_node is not None else ""
                    
                    # Parse category tags
                    categories = [cat.text.strip() for cat in item.findall("category") if cat.text]
                    
                    dt = parse_pub_date(pub_date_str)
                    if dt and dt >= cutoff_date:
                        category = categorize_article(title, description, categories)
                        all_articles.append({
                            "source": source_name,
                            "title": title,
                            "link": link,
                            "pub_date": dt,
                            "description": description[:200] + "..." if len(description) > 200 else description,
                            "category": category
                        })
                        source_count += 1
                print(f"  Found {source_count} articles from past {days_back} days.")
        except Exception as e:
            print(f"  Error fetching from {source_name}: {e}")
            
    return all_articles

def add_hyperlink(paragraph, url, text, color="1F4D78", size_pt=11, bold=True):
    part = paragraph.part
    from docx.opc.constants import RELATIONSHIP_TYPE
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:rId'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    rPr.append(rFonts)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(sz)

    new_run.append(rPr)
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def translate_text(text, sl='en', tl='zh-TW'):
    if not text:
        return ""
    import urllib.parse
    import json
    import time
    url = "https://translate.googleapis.com/translate_a/single"
    params = {
        "client": "gtx",
        "sl": sl,
        "tl": tl,
        "dt": "t",
        "q": text
    }
    try:
        full_url = url + "?" + urllib.parse.urlencode(params)
        req = urllib.request.Request(full_url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=10) as resp:
            data = json.loads(resp.read().decode("utf-8"))
            translations = []
            for item in data[0]:
                if item[0]:
                    translations.append(item[0])
            time.sleep(0.05)  # Politeness delay
            return "".join(translations)
    except Exception as e:
        print(f"Translation error: {e}")
        return text

def generate_report(articles, days_back=7):
    # Group by category
    categories_order = [
        "游泳 (Swimming)",
        "騎車 (Cycling)",
        "跑步 (Running)",
        "補給與恢復 (Fueling & Recovery)"
    ]
    
    grouped = {cat: [] for cat in categories_order}
    for art in articles:
        cat = art["category"]
        if cat not in grouped:
            grouped[cat] = []
        grouped[cat].append(art)
        
    today = datetime.now()
    monday = today - timedelta(days=today.weekday())
    week_num = monday.isocalendar().week
    
    # Ensure output dir exists
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    
    # --- 1. Generate English Markdown Report ---
    output_filename_md = f"{today.year}-W{week_num:02d}_當週鐵人新知與文章整理.md"
    output_filepath_md = OUTPUT_DIR / output_filename_md
    
    with open(output_filepath_md, "w", encoding="utf-8") as f:
        f.write(f"# 鐵人三項當週新知與文章整理 (W{week_num:02d})\n\n")
        f.write(f"本報告彙整了過去 {days_back} 天內，全球權威鐵人三項與運動科學網站的最新發表文章。\n")
        f.write(f"**生成時間**：{today.strftime('%Y-%m-%d %H:%M:%S')} | **來源網站**：Slowtwitch, Triathlete, 220 Triathlon, TrainingPeaks, Joe Friel\n\n")
        
        f.write("## 目錄\n")
        for cat in categories_order:
            count = len(grouped[cat])
            f.write(f"- [{cat}](#{cat.lower().replace(' ', '-').replace('(', '').replace(')', '')}) ({count} 篇)\n")
        f.write("\n---\n\n")
        
        for cat in categories_order:
            anchor_name = cat.lower().replace(' ', '-').replace('(', '').replace(')', '')
            f.write(f"## <a name=\"{anchor_name}\"></a>{cat}\n\n")
            if not grouped[cat]:
                f.write("*本週暫無此分類之最新文章.*\n\n")
                continue
                
            grouped[cat].sort(key=lambda x: x["pub_date"], reverse=True)
            for art in grouped[cat]:
                local_time = art["pub_date"].astimezone().strftime("%Y-%m-%d")
                f.write(f"### 🔗 [{art['title']}]({art['link']})\n")
                f.write(f"**來源**: {art['source']} | **日期**: {local_time}\n\n")
                if art["description"]:
                    f.write(f"> {art['description']}\n\n")
                f.write("---\n\n")
                
    print(f"Markdown report generated successfully: {output_filepath_md}")

    # --- 2. Generate English Word (DOCX) Report ---
    output_filename_docx = f"{today.year}-W{week_num:02d}_當週鐵人新知與文章整理.docx"
    output_filepath_docx = OUTPUT_DIR / output_filename_docx
    
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, attr, Inches(1))
    
    # Styles Setup
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    
    def apply_font(run, size=11, bold=False, color="000000"):
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)
        rpr = run._element.get_or_add_rPr()
        fonts = rpr.rFonts
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            rpr.append(fonts)
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        fonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

    # Document Header & Footer
    header = section.header.paragraphs[0]
    header.text = "全球鐵人三項與運動科學最新文章彙整"
    apply_font(header.runs[0], 9, False, "666666")
    
    footer = section.footer.paragraphs[0]
    footer.text = f"W{week_num:02d} 鐵人新知整理報告 | 彙整日期：{today:%Y-%m-%d}"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    apply_font(footer.runs[0], 9, False, "777777")

    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run(f"鐵人三項當週新知與文章整理 (W{week_num:02d})")
    apply_font(run_title, 22, True, "1F4D78")
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(18)
    run_sub = p_sub.add_run(f"彙整時間：{today.strftime('%Y-%m-%d %H:%M:%S')}　｜　時間範圍：過去 {days_back} 天\n來源網站：Slowtwitch, Triathlete, 220 Triathlon, TrainingPeaks, Joe Friel")
    apply_font(run_sub, 9.5, False, "555555")
    
    for cat in categories_order:
        p_cat = doc.add_paragraph()
        p_cat.paragraph_format.space_before = Pt(16)
        p_cat.paragraph_format.space_after = Pt(8)
        p_cat.paragraph_format.keep_with_next = True
        run_cat = p_cat.add_run(cat)
        apply_font(run_cat, 15, True, "1F4D78")
        
        if not grouped[cat]:
            p_none = doc.add_paragraph()
            p_none.paragraph_format.left_indent = Inches(0.2)
            run_none = p_none.add_run("* 本週暫無此分類之最新文章 *")
            apply_font(run_none, 10, False, "888888")
            run_none.italic = True
            continue
            
        grouped[cat].sort(key=lambda x: x["pub_date"], reverse=True)
        
        for art in grouped[cat]:
            p_art = doc.add_paragraph()
            p_art.paragraph_format.space_before = Pt(8)
            p_art.paragraph_format.space_after = Pt(2)
            p_art.paragraph_format.keep_with_next = True
            
            local_time = art["pub_date"].astimezone().strftime("%Y-%m-%d")
            run_meta = p_art.add_run(f"【{art['source']} | {local_time}】 ")
            apply_font(run_meta, 9.5, True, "595959")
            
            add_hyperlink(p_art, art["link"], art["title"], color="1F4D78", size_pt=11, bold=True)
            
            if art["description"]:
                p_desc = doc.add_paragraph()
                p_desc.paragraph_format.left_indent = Inches(0.25)
                p_desc.paragraph_format.space_after = Pt(10)
                run_desc = p_desc.add_run(art["description"])
                apply_font(run_desc, 9.5, False, "444444")
                
    doc.save(output_filepath_docx)
    print(f"DOCX report generated successfully: {output_filepath_docx}")

    # --- 3. Translate to Chinese & Generate Chinese Reports ---
    print("Translating articles to Traditional Chinese...")
    translated_grouped = {cat: [] for cat in categories_order}
    for cat in categories_order:
        for art in grouped[cat]:
            title_zh = translate_text(art["title"])
            desc_zh = translate_text(art["description"]) if art["description"] else ""
            translated_grouped[cat].append({
                "title_zh": title_zh,
                "title_en": art["title"],
                "link": art["link"],
                "source": art["source"],
                "pub_date": art["pub_date"],
                "description_zh": desc_zh
            })

    # Helper function to generate sport summary
    def generate_sport_summary_text(translated_grouped):
        import os
        api_key = os.environ.get("GEMINI_API_KEY")
        key_file = ROOT_DIR / "config" / "gemini_api_key.txt"
        if not api_key and key_file.exists():
            try:
                api_key = key_file.read_text(encoding="utf-8").strip()
            except:
                pass
                
        if api_key:
            print("Using Gemini API to generate sport summary...")
            simplified_data = {}
            for cat_name, arts in translated_grouped.items():
                simplified_data[cat_name] = [
                    {"title": a["title_zh"], "desc": a["description_zh"][:120]} for a in arts[:5]
                ]
            prompt = f"""
你是一位專業的鐵人三項教練與運動科學專家。以下是本週從國際權威網站（如 Slowtwitch, Triathlete, 220 Triathlon 等）抓取並翻譯好的鐵人三項新知文章列表。
請針對「游泳」、「騎車」、「跑步」以及「補給與恢復」這四個類別，分別撰寫一段 100-150 字的「綜合整理重點」，指出本週的新趨勢、關鍵技術或實用建議。
請使用繁體中文撰寫，語氣專業且實用，並直接輸出 Markdown 格式的內容，包含標題「## 五、 當週各項運動新知綜合整理重點」以及各項目的子標題。

文章列表：
{json.dumps(simplified_data, ensure_ascii=False, indent=2)}
"""
            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
            headers = {"Content-Type": "application/json"}
            payload = {
                "contents": [{
                    "parts": [{"text": prompt}]
                }]
            }
            try:
                import urllib.request
                import json
                req = urllib.request.Request(url, data=json.dumps(payload).encode("utf-8"), headers=headers, method="POST")
                with urllib.request.urlopen(req, timeout=30) as resp:
                    res_data = json.loads(resp.read().decode("utf-8"))
                    gpt_text = res_data["candidates"][0]["content"]["parts"][0]["text"]
                    if gpt_text:
                        return gpt_text
            except Exception as e:
                print("Gemini API summarization failed, falling back to extractive summary:", e)
                
        print("Generating fallback extractive summary...")
        lines = []
        lines.append("## 五、 當週各項運動新知綜合整理重點\n\n")
        for cat_name in ["游泳 (Swimming)", "騎車 (Cycling)", "跑步 (Running)", "補給與恢復 (Fueling & Recovery)"]:
            lines.append(f"### {cat_name}核心要點\n")
            arts = translated_grouped[cat_name]
            if not arts:
                lines.append("- 本週無此項目新知。\n\n")
                continue
            for a in arts[:3]:
                desc_clean = a["description_zh"].replace("...", "").strip()
                lines.append(f"* **{a['title_zh']}**：{desc_clean}\n")
            lines.append("\n")
        return "".join(lines)

    summary_text = generate_sport_summary_text(translated_grouped)

    # Generate Chinese Markdown
    output_filename_md_zh = f"{today.year}-W{week_num:02d}_當週鐵人新知與文章整理_中文版.md"
    output_filepath_md_zh = OUTPUT_DIR / output_filename_md_zh
    
    with open(output_filepath_md_zh, "w", encoding="utf-8") as f:
        f.write(f"# 鐵人三項當週新知與文章整理 (W{week_num:02d}) - 中文翻譯版\n\n")
        f.write(f"本報告彙整了過去 {days_back} 天內，全球權威鐵人三項與運動科學網站的最新發表文章並翻譯為繁體中文。\n")
        f.write(f"**生成時間**：{today.strftime('%Y-%m-%d %H:%M:%S')} | **來源網站**：Slowtwitch, Triathlete, 220 Triathlon, TrainingPeaks, Joe Friel\n\n")
        
        f.write("## 目錄\n")
        for cat in categories_order:
            count = len(translated_grouped[cat])
            f.write(f"- [{cat}](#{cat.lower().replace(' ', '-').replace('(', '').replace(')', '')}) ({count} 篇)\n")
        f.write("\n---\n\n")
        
        for cat in categories_order:
            anchor_name = cat.lower().replace(' ', '-').replace('(', '').replace(')', '')
            f.write(f"## <a name=\"{anchor_name}\"></a>{cat}\n\n")
            if not translated_grouped[cat]:
                f.write("*本週暫無此分類之最新文章.*\n\n")
                continue
                
            translated_grouped[cat].sort(key=lambda x: x["pub_date"], reverse=True)
            for art in translated_grouped[cat]:
                local_time = art["pub_date"].astimezone().strftime("%Y-%m-%d")
                f.write(f"### 🔗 [{art['title_zh']}]({art['link']})\n")
                f.write(f"**英文原名**: *{art['title_en']}*\n\n")
                f.write(f"**來源**: {art['source']} | **日期**: {local_time}\n\n")
                if art["description_zh"]:
                    f.write(f"> {art['description_zh']}\n\n")
                f.write("---\n\n")
        
        f.write("\n---\n\n")
        f.write(summary_text)
                
    print(f"Chinese Markdown report generated successfully: {output_filepath_md_zh}")

    # Generate Chinese Word (DOCX)
    output_filename_docx_zh = f"{today.year}-W{week_num:02d}_當週鐵人新知與文章整理_中文版.docx"
    output_filepath_docx_zh = OUTPUT_DIR / output_filename_docx_zh
    
    doc_zh = Document()
    
    # Page setup
    section_zh = doc_zh.sections[0]
    section_zh.page_width = Inches(8.5)
    section_zh.page_height = Inches(11)
    for attr in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section_zh, attr, Inches(1))
        
    # Styles Setup
    normal_zh = doc_zh.styles["Normal"]
    normal_zh.font.name = "Calibri"
    normal_zh.font.size = Pt(11)
    normal_zh.paragraph_format.space_after = Pt(6)
    normal_zh.paragraph_format.line_spacing = 1.15

    # Document Header & Footer
    header_zh = section_zh.header.paragraphs[0]
    header_zh.text = "全球鐵人三項與運動科學最新文章彙整 - 中文翻譯版"
    apply_font(header_zh.runs[0], 9, False, "666666")
    
    footer_zh = section_zh.footer.paragraphs[0]
    footer_zh.text = f"W{week_num:02d} 鐵人新知整理報告 (中文版) | 彙整日期：{today:%Y-%m-%d}"
    footer_zh.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    apply_font(footer_zh.runs[0], 9, False, "777777")

    # Title
    p_title_zh = doc_zh.add_paragraph()
    p_title_zh.paragraph_format.space_before = Pt(0)
    p_title_zh.paragraph_format.space_after = Pt(4)
    run_title_zh = p_title_zh.add_run(f"鐵人三項當週新知與文章整理 (W{week_num:02d}) - 中文版")
    apply_font(run_title_zh, 22, True, "1F4D78")
    
    # Subtitle
    p_sub_zh = doc_zh.add_paragraph()
    p_sub_zh.paragraph_format.space_after = Pt(18)
    run_sub_zh = p_sub_zh.add_run(f"彙整時間：{today.strftime('%Y-%m-%d %H:%M:%S')}　｜　時間範圍：過去 {days_back} 天\n來源網站：Slowtwitch, Triathlete, 220 Triathlon, TrainingPeaks, Joe Friel (已翻譯為繁體中文)")
    apply_font(run_sub_zh, 9.5, False, "555555")
    
    for cat in categories_order:
        p_cat = doc_zh.add_paragraph()
        p_cat.paragraph_format.space_before = Pt(16)
        p_cat.paragraph_format.space_after = Pt(8)
        p_cat.paragraph_format.keep_with_next = True
        run_cat = p_cat.add_run(cat)
        apply_font(run_cat, 15, True, "1F4D78")
        
        if not translated_grouped[cat]:
            p_none = doc_zh.add_paragraph()
            p_none.paragraph_format.left_indent = Inches(0.2)
            run_none = p_none.add_run("* 本週暫無此分類之最新文章 *")
            apply_font(run_none, 10, False, "888888")
            run_none.italic = True
            continue
            
        translated_grouped[cat].sort(key=lambda x: x["pub_date"], reverse=True)
        
        for art in translated_grouped[cat]:
            p_art = doc_zh.add_paragraph()
            p_art.paragraph_format.space_before = Pt(8)
            p_art.paragraph_format.space_after = Pt(2)
            p_art.paragraph_format.keep_with_next = True
            
            local_time = art["pub_date"].astimezone().strftime("%Y-%m-%d")
            run_meta = p_art.add_run(f"【{art['source']} | {local_time}】 ")
            apply_font(run_meta, 9.5, True, "595959")
            
            add_hyperlink(p_art, art["link"], art["title_zh"], color="1F4D78", size_pt=11, bold=True)
            
            p_eng = doc_zh.add_paragraph()
            p_eng.paragraph_format.left_indent = Inches(0.25)
            p_eng.paragraph_format.space_after = Pt(2)
            run_eng = p_eng.add_run(f"原名: {art['title_en']}")
            apply_font(run_eng, 8.5, False, "777777")
            run_eng.italic = True
            
            if art["description_zh"]:
                p_desc = doc_zh.add_paragraph()
                p_desc.paragraph_format.left_indent = Inches(0.25)
                p_desc.paragraph_format.space_after = Pt(10)
                run_desc = p_desc.add_run(art["description_zh"])
                apply_font(run_desc, 9.5, False, "444444")

    # Helper function to append markdown summary to docx
    def append_markdown_summary_to_docx(doc, summary_text):
        lines = summary_text.replace("\r\n", "\n").split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith("## "):
                text = line.replace("## ", "").strip()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(20)
                p.paragraph_format.space_after = Pt(10)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(text)
                apply_font(run, 15, True, "1F4D78")
            elif line.startswith("### ") or line.startswith("#### "):
                text = line.replace("### ", "").replace("#### ", "").strip()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(text)
                apply_font(run, 12, True, "2E74B5")
            elif line.startswith("* ") or line.startswith("- "):
                text = line[2:].strip()
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Inches(0.25)
                
                if "**" in text:
                    parts = text.split("**")
                    is_bold = False
                    for part in parts:
                        if not part:
                            is_bold = not is_bold
                            continue
                        run = p.add_run(part)
                        apply_font(run, 10, is_bold, "333333")
                        is_bold = not is_bold
                else:
                    run = p.add_run(text)
                    apply_font(run, 10, False, "333333")
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(line)
                apply_font(run, 10, False, "333333")

    append_markdown_summary_to_docx(doc_zh, summary_text)
                
    doc_zh.save(output_filepath_docx_zh)
    print(f"Chinese DOCX report generated successfully: {output_filepath_docx_zh}")
    
    return output_filepath_docx

def main():
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--days", type=int, default=7, help="Number of days to look back")
    args = parser.parse_args()
    
    articles = fetch_recent_articles(days_back=args.days)
    if articles:
        report_path = generate_report(articles, days_back=args.days)
        print(f"SUCCESS:{report_path}")
    else:
        print("NO_NEW_ARTICLES")

if __name__ == "__main__":
    main()
