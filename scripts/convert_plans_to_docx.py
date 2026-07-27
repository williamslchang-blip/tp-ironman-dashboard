import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_full_plan_docx():
    md_path = r"C:\Users\User\Desktop\TP\outputs\12_week_ironman_full_training_plan.md"
    docx_path = r"C:\Users\User\Desktop\TP\outputs\12_week_ironman_full_training_plan.docx"
    
    if not os.path.exists(md_path):
        print("MD file not found")
        return

    doc = docx.Document()
    
    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.splitlines()
    in_table = False
    table_lines = []

    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        rows = [t.strip("|").split("|") for t in table_lines if t.strip()]
        if len(rows) >= 2:
            # Check if second line is separator
            if "---" in rows[1][0] or ":---" in rows[1][0]:
                headers = [c.strip() for c in rows[0]]
                data_rows = [[c.strip() for c in r] for r in rows[2:]]
            else:
                headers = [c.strip() for c in rows[0]]
                data_rows = [[c.strip() for c in r] for r in rows[1:]]

            tbl = doc.add_table(rows=len(data_rows)+1, cols=len(headers))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header row
            hdr_cells = tbl.rows[0].cells
            for idx, text in enumerate(headers):
                hdr_cells[idx].text = text
                set_cell_background(hdr_cells[idx], "1A365D")  # Dark blue
                p = hdr_cells[idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.name = "微軟正黑體"
                    run.font.size = Pt(10)

            # Data rows
            for r_idx, r_data in enumerate(data_rows):
                row_cells = tbl.rows[r_idx+1].cells
                bg_color = "F7FAFC" if r_idx % 2 == 1 else "FFFFFF"
                for c_idx, val in enumerate(r_data):
                    if c_idx < len(row_cells):
                        row_cells[c_idx].text = val
                        set_cell_background(row_cells[c_idx], bg_color)
                        p = row_cells[c_idx].paragraphs[0]
                        for run in p.runs:
                            run.font.name = "微軟正黑體"
                            run.font.size = Pt(9.5)
                        if c_idx == 0:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph() # space after table
        table_lines = []

    for line in lines:
        sline = line.strip()
        if sline.startswith("|"):
            in_table = True
            table_lines.append(sline)
            continue
        elif in_table:
            in_table = False
            flush_table()

        if sline.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(sline[2:].strip())
            run.font.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(26, 54, 93) # Dark Blue
            run.font.name = "微軟正黑體"
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
        elif sline.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(sline[3:].strip())
            run.font.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(43, 108, 176) # Blue
            run.font.name = "微軟正黑體"
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
        elif sline.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(sline[4:].strip())
            run.font.bold = True
            run.font.size = Pt(12.5)
            run.font.color.rgb = RGBColor(45, 55, 72) # Slate
            run.font.name = "微軟正黑體"
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif sline.startswith("- ") or sline.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            text = sline[2:].strip()
            # Handle inline bolding
            parts = text.split("**")
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = "微軟正黑體"
                run.font.size = Pt(10.5)
                if idx % 2 == 1:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(26, 54, 93)
            p.paragraph_format.space_after = Pt(3)
        elif sline.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(sline[2:].strip())
            run.font.italic = True
            run.font.color.rgb = RGBColor(113, 128, 150)
            run.font.name = "微軟正黑體"
            run.font.size = Pt(10)
        elif sline == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
        elif sline:
            p = doc.add_paragraph()
            parts = sline.split("**")
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = "微軟正黑體"
                run.font.size = Pt(10.5)
                if idx % 2 == 1:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(26, 54, 93)
            p.paragraph_format.space_after = Pt(4)

    if in_table:
        flush_table()

    doc.save(docx_path)
    print(f"Saved {docx_path}")

def create_daily_plan_docx():
    md_path = r"C:\Users\User\Desktop\TP\outputs\12_week_ironman_daily_training_plan.md"
    docx_path = r"C:\Users\User\Desktop\TP\outputs\12_week_ironman_daily_training_plan.docx"

    if not os.path.exists(md_path):
        print("MD file not found")
        return

    doc = docx.Document()
    
    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.splitlines()

    for line in lines:
        sline = line.strip()
        if sline.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(sline[2:].strip())
            run.font.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(26, 54, 93)
            run.font.name = "微軟正黑體"
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
        elif sline.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(sline[3:].strip())
            run.font.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(43, 108, 176)
            run.font.name = "微軟正黑體"
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
        elif sline.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(sline[4:].strip())
            run.font.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(45, 55, 72)
            run.font.name = "微軟正黑體"
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
        elif sline.startswith("#### "):
            p = doc.add_paragraph()
            run = p.add_run(sline[5:].strip())
            run.font.bold = True
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(43, 108, 176)
            run.font.name = "微軟正黑體"
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
        elif sline.startswith("- ") or sline.startswith("1. ") or sline.startswith("2. ") or sline.startswith("3. ") or sline.startswith("4. ") or sline.startswith("5. "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            text = sline
            parts = text.split("**")
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = "微軟正黑體"
                run.font.size = Pt(10)
                if idx % 2 == 1:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(26, 54, 93)
            p.paragraph_format.space_after = Pt(2)
        elif sline.startswith("  "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            parts = sline.strip().split("**")
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = "微軟正黑體"
                run.font.size = Pt(9.5)
                if idx % 2 == 1:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(43, 108, 176)
            p.paragraph_format.space_after = Pt(2)
        elif sline == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        elif sline:
            p = doc.add_paragraph()
            parts = sline.split("**")
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = "微軟正黑體"
                run.font.size = Pt(10)
                if idx % 2 == 1:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(26, 54, 93)
            p.paragraph_format.space_after = Pt(4)

    doc.save(docx_path)
    print(f"Saved {docx_path}")

def create_consolidated_plan_docx():
    md_path = r"C:\Users\User\Desktop\TP\outputs\12_week_ironman_consolidated_master_plan.md"
    docx_path = r"C:\Users\User\Desktop\TP\outputs\12_week_ironman_custom_master_plan.docx"


    if not os.path.exists(md_path):
        print("MD file not found")
        return

    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.splitlines()
    in_table = False
    table_lines = []

    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        rows = [t.strip("|").split("|") for t in table_lines if t.strip()]
        if len(rows) >= 2:
            if "---" in rows[1][0] or ":---" in rows[1][0]:
                headers = [c.strip() for c in rows[0]]
                data_rows = [[c.strip() for c in r] for r in rows[2:]]
            else:
                headers = [c.strip() for c in rows[0]]
                data_rows = [[c.strip() for c in r] for r in rows[1:]]

            tbl = doc.add_table(rows=len(data_rows)+1, cols=len(headers))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            hdr_cells = tbl.rows[0].cells
            for idx, text in enumerate(headers):
                hdr_cells[idx].text = text
                set_cell_background(hdr_cells[idx], "1A365D")
                p = hdr_cells[idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.name = "微軟正黑體"
                    run.font.size = Pt(9.5)

            for r_idx, r_data in enumerate(data_rows):
                row_cells = tbl.rows[r_idx+1].cells
                bg_color = "F7FAFC" if r_idx % 2 == 1 else "FFFFFF"
                for c_idx, val in enumerate(r_data):
                    if c_idx < len(row_cells):
                        row_cells[c_idx].text = val
                        set_cell_background(row_cells[c_idx], bg_color)
                        p = row_cells[c_idx].paragraphs[0]
                        for run in p.runs:
                            run.font.name = "微軟正黑體"
                            run.font.size = Pt(9)
                        if c_idx == 0:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        table_lines = []

    for line in lines:
        sline = line.strip()
        if sline.startswith("|"):
            in_table = True
            table_lines.append(sline)
            continue
        elif in_table:
            in_table = False
            flush_table()

        if sline.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(sline[2:].strip())
            run.font.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(26, 54, 93)
            run.font.name = "微軟正黑體"
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
        elif sline.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(sline[3:].strip())
            run.font.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(43, 108, 176)
            run.font.name = "微軟正黑體"
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
        elif sline.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(sline[4:].strip())
            run.font.bold = True
            run.font.size = Pt(12.5)
            run.font.color.rgb = RGBColor(45, 55, 72)
            run.font.name = "微軟正黑體"
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif sline.startswith("#### "):
            p = doc.add_paragraph()
            run = p.add_run(sline[5:].strip())
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(43, 108, 176)
            run.font.name = "微軟正黑體"
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
        elif sline.startswith("- ") or sline.startswith("1. ") or sline.startswith("2. ") or sline.startswith("3. ") or sline.startswith("4. ") or sline.startswith("5. "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            parts = sline.split("**")
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = "微軟正黑體"
                run.font.size = Pt(10)
                if idx % 2 == 1:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(26, 54, 93)
            p.paragraph_format.space_after = Pt(2)
        elif sline.startswith("  "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            parts = sline.strip().split("**")
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = "微軟正黑體"
                run.font.size = Pt(9.5)
                if idx % 2 == 1:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(43, 108, 176)
            p.paragraph_format.space_after = Pt(2)
        elif sline == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        elif sline:
            p = doc.add_paragraph()
            parts = sline.split("**")
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = "微軟正黑體"
                run.font.size = Pt(10)
                if idx % 2 == 1:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(26, 54, 93)
            p.paragraph_format.space_after = Pt(4)

    if in_table:
        flush_table()

    try:
        doc.save(docx_path)
        print(f"Saved {docx_path}")
    except PermissionError:
        fallback_path = docx_path.replace(".docx", "_custom.docx")
        doc.save(fallback_path)
        print(f"Saved fallback {fallback_path}")

if __name__ == "__main__":
    create_full_plan_docx()
    create_daily_plan_docx()
    create_consolidated_plan_docx()


