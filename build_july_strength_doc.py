from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(r"C:\Users\User\Desktop\TP")
OUT_FILE = OUT_DIR / "2026_07_strength_schedule.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="Arial", size=11, bold=False, color="000000"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)


def style_paragraph(paragraph, *, before=0, after=8, line=1.15, align=WD_ALIGN_PARAGRAPH.LEFT):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    paragraph.alignment = align


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
for attr in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
    setattr(section, attr, Inches(1))
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0, 0, 0)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15

for name, size, before, after, color in [
    ("Heading 1", 20, 20, 6, "000000"),
    ("Heading 2", 16, 18, 6, "000000"),
    ("Heading 3", 14, 16, 4, "434343"),
]:
    st = styles[name]
    st.font.name = "Arial"
    st.font.size = Pt(size)
    st.font.bold = False
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.line_spacing = 1.15

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(3)
title.paragraph_format.line_spacing = 1.0
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run("2026 年 7 月肌力課表")
set_run_font(run, size=26, bold=False)

subtitle = doc.add_paragraph()
style_paragraph(subtitle, before=0, after=8, line=1.15)
subtitle_run = subtitle.add_run("依 TP 固定週課表排入：週一休息後主課、週三/週五短課、週日恢復短課")
set_run_font(subtitle_run, size=11, color="555555")

note = doc.add_paragraph()
style_paragraph(note, before=0, after=10, line=1.15)
note_run = note.add_run("原則：主課安排在週一晚間，短課安排在游泳日後的晚間或訓練後 30 到 60 分鐘內；若當天腿部疲勞高，先保留短課或取消。")
set_run_font(note_run, size=11)

headers = ["日期", "TP 課表", "肌力時段與安排"]
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.autofit = False
table.style = "Table Grid"
widths = [Inches(1.0), Inches(2.45), Inches(3.05)]
hdr_cells = table.rows[0].cells
for cell, text, width in zip(hdr_cells, headers, widths):
    cell.width = width
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "F2F4F7")
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=11, bold=True)
set_repeat_table_header(table.rows[0])

weekly_tp = {
    0: "休息",
    1: "間歇 + 配速跑",
    2: "游泳",
    3: "單車訓練台",
    4: "游泳",
    5: "長騎 + 跑步",
    6: "跑步 + 游泳",
}

strength_plan = {
    0: "主課 30 到 40 分鐘，建議晚上 19:00 到 20:30；若整週疲勞偏高，可縮成 20 分鐘核心/小腿。",
    1: "不排肌力。若想加，最多做 10 到 15 分鐘活動度與放鬆，避免影響間歇品質。",
    2: "短課 15 到 20 分鐘，建議晚間 19:30 到 21:00，或游泳後 30 到 60 分鐘內完成。",
    3: "不排肌力。保留單車台品質。",
    4: "短課 15 到 20 分鐘，建議晚間 19:30 到 21:00，做恢復型核心與髖穩定。",
    5: "不排肌力。長騎 + 跑步日只做活動度，避免干擾週末主課。",
    6: "恢復短課 15 分鐘，建議晚間 18:00 到 20:00；若週六長騎跑很硬，這天可取消。",
}

weekday_tw = ["週一", "週二", "週三", "週四", "週五", "週六", "週日"]

for day in range(1, 32):
    dt = date(2026, 7, day)
    row = table.add_row()
    row.height = None
    cells = row.cells
    for idx, cell in enumerate(cells):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        if idx == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"{day:02d}\n{weekday_tw[dt.weekday()]}")
            set_run_font(r, size=10)
        elif idx == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(weekly_tp[dt.weekday()])
            set_run_font(r, size=10)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            muscle = strength_plan[dt.weekday()]
            if dt.weekday() == 0:
                text = f"{muscle}\n內容：下肢主課 A。"
            elif dt.weekday() in (2, 4):
                text = f"{muscle}\n內容：短課，見下方短課模板。"
            elif dt.weekday() == 6:
                text = f"{muscle}\n內容：恢復短課；若隔天恢復差，直接取消。"
            else:
                text = muscle
            lines = text.split("\n")
            for i, line in enumerate(lines):
                if i > 0:
                    p.add_run().add_break()
                r = p.add_run(line)
                set_run_font(r, size=10)

doc.add_paragraph()

h2 = doc.add_paragraph(style="Heading 2")
h2_run = h2.add_run("主課與短課內容")
set_run_font(h2_run, size=16, bold=False)

main_intro = doc.add_paragraph()
style_paragraph(main_intro, before=0, after=6, line=1.15)
r = main_intro.add_run("主課放在週一，重點是下肢力量、後鏈與跑步抗疲勞能力。")
set_run_font(r, size=11)

main_a = doc.add_paragraph(style="Heading 3")
main_a_run = main_a.add_run("主課 A")
set_run_font(main_a_run, size=14, bold=False, color="434343")

for text in [
    "熱身 8 分鐘：髖關節活動、徒手深蹲、臀橋、弓箭步、提踵。",
    "分腿蹲 3 組 x 6 到 8 下/邊。",
    "羅馬尼亞硬舉 3 組 x 6 到 8 下。",
    "單腳硬舉 2 組 x 6 到 8 下/邊。",
    "側棒式 2 組 x 30 到 45 秒/邊。",
    "Dead bug 2 組 x 8 到 10 下/邊。",
    "提踵 3 組 x 12 到 15 下。",
]:
    p = doc.add_paragraph(style="List Bullet")
    style_paragraph(p, before=0, after=3, line=1.15)
    r = p.add_run(text)
    set_run_font(r, size=11)

main_b = doc.add_paragraph(style="Heading 3")
main_b_run = main_b.add_run("主課 B")
set_run_font(main_b_run, size=14, bold=False, color="434343")

for text in [
    "熱身 8 分鐘：臀橋、髖外展、單腳平衡、徒手弓箭步、提踵。",
    "臀推或橋式 3 組 x 8 到 10 下。",
    "Step-up 3 組 x 6 到 8 下/邊。",
    "Pallof press 2 組 x 10 到 12 下/邊。",
    "側向走或怪獸走 2 組 x 10 到 15 步/邊。",
    "離心提踵 3 組 x 8 到 12 下。",
    "髖屈肌伸展 2 組 x 30 秒/邊。",
]:
    p = doc.add_paragraph(style="List Bullet")
    style_paragraph(p, before=0, after=3, line=1.15)
    r = p.add_run(text)
    set_run_font(r, size=11)

short_h = doc.add_paragraph(style="Heading 3")
short_h_run = short_h.add_run("短課")
set_run_font(short_h_run, size=14, bold=False, color="434343")

for text in [
    "死蟲 2 組 x 8 到 10 下/邊。",
    "側棒式 2 組 x 30 秒/邊。",
    "Bird dog 2 組 x 8 下/邊。",
    "彈力帶側走 2 組 x 10 到 15 步/邊。",
    "單腳提踵 2 組 x 12 到 15 下/邊。",
    "小腿伸展 2 組 x 30 秒/邊。",
]:
    p = doc.add_paragraph(style="List Bullet")
    style_paragraph(p, before=0, after=3, line=1.15)
    r = p.add_run(text)
    set_run_font(r, size=11)

foot = doc.add_paragraph()
style_paragraph(foot, before=6, after=0, line=1.15)
foot_r = foot.add_run("執行規則：主課保留 2 到 3 下餘裕，不做到力竭；若當天腿部疲勞高，優先保留短課或直接取消，不硬補。")
set_run_font(foot_r, size=11)

doc.save(OUT_FILE)
print(str(OUT_FILE))
